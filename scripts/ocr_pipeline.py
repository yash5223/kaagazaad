import cv2
import numpy as np
import pytesseract
import re
import random
import argparse
import concurrent.futures
from difflib import SequenceMatcher
try:
    import joblib
    from sklearn.linear_model import LogisticRegression
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.pipeline import Pipeline
    _ML_AVAILABLE = True
except ImportError:
    _ML_AVAILABLE = False
import json
import sys
import os
from datetime import datetime, timezone
from statistics import mean
def preprocess_image(path, max_dim=2000):
    img = cv2.imread(path)
    if img is None:
        raise FileNotFoundError(f"Could not read image: {path}")
    return preprocess_image_array(img, max_dim=max_dim)
def preprocess_image_array(img, max_dim=2000):
    h, w = img.shape[:2]
    scale = max_dim / max(h, w)
    if scale != 1.0:
        img = cv2.resize(img, (int(w * scale), int(h * scale)),
                         interpolation=cv2.INTER_CUBIC if scale > 1 else cv2.INTER_AREA)
    gray = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
    gray = clahe.apply(gray)
    denoised = cv2.fastNlMeansDenoising(gray, h=10)
    deskewed = _reorient_with_osd(denoised)
    thresh = cv2.adaptiveThreshold(
        deskewed, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C,
        cv2.THRESH_BINARY, 31, 15
    )
    return {"original": img, "processed": thresh, "gray_deskewed": deskewed}
def _reorient_with_osd(gray_img):
    try:
        osd = pytesseract.image_to_osd(gray_img, output_type=pytesseract.Output.DICT)
        rotate = osd.get("rotate", 0)
    except Exception:
        return gray_img
    if not rotate:
        return gray_img
    (h, w) = gray_img.shape
    center = (w // 2, h // 2)
    M = cv2.getRotationMatrix2D(center, -rotate, 1.0)
    if rotate in (90, 270):
        M[0, 2] += (h - w) / 2
        M[1, 2] += (w - h) / 2
        return cv2.warpAffine(gray_img, M, (h, w), flags=cv2.INTER_CUBIC,
                            borderMode=cv2.BORDER_REPLICATE)
    return cv2.warpAffine(gray_img, M, (w, h), flags=cv2.INTER_CUBIC,
                            borderMode=cv2.BORDER_REPLICATE)
def _run_tesseract(img, lang, psm):
    config = f"--oem 3 --psm {psm}"
    data = pytesseract.image_to_data(
        img, lang=lang, config=config, output_type=pytesseract.Output.DICT
    )
    words = []
    for i in range(len(data["text"])):
        text = data["text"][i].strip()
        conf = float(data["conf"][i])
        if text and conf >= 0:
            words.append({
                "text": text, "conf": conf,
                "left": data["left"][i], "top": data["top"][i],
                "width": data["width"][i], "height": data["height"][i],
                "line_num": data["line_num"][i], "block_num": data["block_num"][i],
                "par_num": data["par_num"][i],
            })
    raw_text = pytesseract.image_to_string(img, lang=lang, config=config).strip()
    avg_conf = mean([w["conf"] for w in words]) if words else 0.0
    return {"words": words, "raw_text": raw_text, "avg_confidence": round(avg_conf, 2)}
_OCR_MODEL_CONFIGS = [
    ("thresholded_psm6", "processed", 6),
    ("grayscale_psm6", "gray_deskewed", 6),
    ("grayscale_psm4", "gray_deskewed", 4),
    ("grayscale_psm11", "gray_deskewed", 11),
]
# Adaptive fast-path: most clean scans/photos are already well handled by the
# first (best-performing) model config. Running all 4 tesseract passes on
# every document is the single biggest cost in the pipeline, so when the
# primary pass comes back clean and confident we skip the other 3 entirely.
# Harder/noisier documents still fall through to the full ensemble vote.
FAST_PATH_MIN_CONFIDENCE = 78.0
FAST_PATH_MIN_WORDS = 8
def _run_single_model(stages, lang, cfg):
    name, stage_key, psm = cfg
    try:
        result = _run_tesseract(stages[stage_key], lang, psm)
    except Exception:
        return None
    result["model"] = name
    return result
def _run_ocr_models(stages, lang="eng", adaptive=True):
    configs = _OCR_MODEL_CONFIGS
    results = []
    remaining = configs
    if adaptive:
        primary = _run_single_model(stages, lang, configs[0])
        if primary is not None:
            results.append(primary)
            if (primary["avg_confidence"] >= FAST_PATH_MIN_CONFIDENCE
                    and len(primary["words"]) >= FAST_PATH_MIN_WORDS):
                return results
        remaining = configs[1:]
    # Run the remaining tesseract passes concurrently. pytesseract shells out
    # to the tesseract binary as a subprocess per call, so this is true
    # parallelism (not limited by the GIL) and scales with available cores.
    if remaining:
        with concurrent.futures.ThreadPoolExecutor(max_workers=min(4, len(remaining))) as pool:
            for result in pool.map(lambda cfg: _run_single_model(stages, lang, cfg), remaining):
                if result is not None:
                    results.append(result)
    return results
def _ensemble_words(candidates):
    if not candidates:
        return {"words": [], "raw_text": "", "avg_confidence": 0.0,
                "ensemble_method": None, "models_compared": []}
    models_compared = [{"model": c["model"], "avg_confidence": c["avg_confidence"]}
                       for c in candidates]
    if len(candidates) == 1:
        only = candidates[0]
        return {
            "words": only["words"], "raw_text": only["raw_text"],
            "avg_confidence": only["avg_confidence"],
            "ensemble_method": "single_pass_fast",
            "models_compared": models_compared, "winning_model": only["model"],
        }
    word_counts = {len(c["words"]) for c in candidates}
    best_overall = max(candidates, key=lambda c: c["avg_confidence"])
    if len(word_counts) != 1 or not best_overall["words"]:
        fallback_layout = analyze_layout(best_overall["words"])
        raw_text = "\n".join(fallback_layout["lines"]) if fallback_layout["lines"] else best_overall["raw_text"]
        return {
            "words": best_overall["words"], "raw_text": raw_text,
            "avg_confidence": best_overall["avg_confidence"],
            "ensemble_method": "best_model_fallback",
            "models_compared": models_compared, "winning_model": best_overall["model"],
        }
    resolved_words = []
    for i in range(len(best_overall["words"])):
        slot = [(c["model"], c["words"][i]) for c in candidates]
        counts = {}
        for _, w in slot:
            counts[w["text"]] = counts.get(w["text"], 0) + 1
        majority_text, majority_n = max(counts.items(), key=lambda kv: kv[1])
        if majority_n > len(slot) / 2:
            agreeing = [(m, w) for m, w in slot if w["text"] == majority_text]
        else:
            agreeing = slot
        winner_model, winner_word = max(agreeing, key=lambda mw: mw[1]["conf"])
        resolved_words.append({**winner_word, "winning_model": winner_model})
    layout = analyze_layout(resolved_words)
    raw_text = "\n".join(layout["lines"])
    avg_conf = mean([w["conf"] for w in resolved_words]) if resolved_words else 0.0
    return {
        "words": resolved_words, "raw_text": raw_text, "avg_confidence": round(avg_conf, 2),
        "ensemble_method": "word_level_vote", "models_compared": models_compared,
        "winning_model": None,
    }
def detect_and_recognize(stages, lang="eng", min_conf=0, adaptive=True):
    candidates = _run_ocr_models(stages, lang=lang, adaptive=adaptive)
    if not candidates:
        return {"words": [], "raw_text": "", "avg_confidence": 0.0,
                "ensemble_method": None, "models_compared": []}
    result = _ensemble_words(candidates)
    if min_conf > 0:
        result["words"] = [w for w in result["words"] if w["conf"] >= min_conf]
    return result
def _print_ocr_model_summary(ocr_result, file=sys.stderr):
    models = ocr_result.get("models_compared")
    if not models:
        return
    print("OCR Model Comparison:", file=file)
    for m in models:
        print(f"  {m['model']}: avg_confidence={m['avg_confidence']}", file=file)
    method = ocr_result.get("ensemble_method")
    if method == "single_pass_fast":
        print("  Result: primary pass was clean/confident enough to skip the "
              "other 3 models (adaptive fast path)", file=file)
    elif method == "word_level_vote":
        print("  Result: word-by-word vote across all models "
              "(majority text, ties broken by confidence)", file=file)
    elif method == "best_model_fallback":
        print(f"  Result: models segmented the page differently, so the "
              f"single most confident model won ({ocr_result.get('winning_model')})",
              file=file)
def analyze_layout(words):
    if not words:
        return {"lines": [], "line_count": 0, "table_detected": False, "table_rows": []}
    words_sorted = sorted(words, key=lambda w: (w["top"], w["left"]))
    rows = []
    for w in words_sorted:
        placed = False
        for row in rows:
            tol = max(row["height"], w["height"]) * 0.6
            if abs(w["top"] - row["top"]) <= tol:
                row["words"].append(w)
                row["height"] = max(row["height"], w["height"])
                placed = True
                break
        if not placed:
            rows.append({"top": w["top"], "height": w["height"], "words": [w]})
    line_list = []
    for row in rows:
        ws_sorted = sorted(row["words"], key=lambda x: x["left"])
        line_text = " ".join(w["text"] for w in ws_sorted)
        top = min(w["top"] for w in ws_sorted)
        line_list.append({"text": line_text, "top": top, "words": ws_sorted})
    line_list.sort(key=lambda l: l["top"])
    table_rows = []
    for line in line_list:
        cols = _split_into_columns(line["words"])
        if len(cols) >= 2:
            table_rows.append(cols)
    table_detected = len(table_rows) >= 2 and len({len(r) for r in table_rows}) <= 2
    return {
        "lines": [l["text"] for l in line_list],
        "line_count": len(line_list),
        "table_detected": table_detected,
        "table_rows": table_rows if table_detected else [],
    }
def _split_into_columns(words_in_line, gap_threshold=40):
    if not words_in_line:
        return []
    cols = [[words_in_line[0]["text"]]]
    prev_right = words_in_line[0]["left"] + words_in_line[0]["width"]
    for w in words_in_line[1:]:
        gap = w["left"] - prev_right
        if gap > gap_threshold:
            cols.append([w["text"]])
        else:
            cols[-1].append(w["text"])
        prev_right = w["left"] + w["width"]
    return [" ".join(c) for c in cols]
FIELD_PATTERNS = {
    "emails": re.compile(r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}"),
    "dates": re.compile(r"\b(\d{1,4}[/-]\d{1,2}[/-]\d{1,4}|\d{1,2}\s+[A-Za-z]{3,9}\s+\d{2,4})\b"),
    "amounts": re.compile(r"(?:[$€£₹]\s?\d[\d,]*\.?\d{0,2})|\b\d[\d,]*\.\d{2}\b"),
    "urls": re.compile(r"https?://\S+|www\.\S+"),
    "aadhaar_numbers": re.compile(r"\b\d{4}\s?\d{4}\s?\d{4}\b"),
    "pan_numbers": re.compile(r"\b[A-Z]{5}[0-9]{4}[A-Z]\b"),
    "gstin_numbers": re.compile(r"\b\d{2}[A-Z]{5}\d{4}[A-Z][A-Z\d]Z[A-Z\d]\b"),
    "ifsc_codes": re.compile(r"\b[A-Z]{4}0[A-Z0-9]{6}\b"),
    "vehicle_reg_numbers": re.compile(r"\b[A-Z]{2}[-\s]?\d{1,2}[-\s]?[A-Z]{1,3}[-\s]?\d{4}\b"),
    "passport_numbers": re.compile(r"\b[A-PR-WYa-pr-wy][0-9]{7}\b"),
    "pincodes": re.compile(r"\b\d{6}\b"),
    "seat_no": re.compile(r"SEAT\s*NO\.?\s*[:\-]?\s*([A-Z0-9\/]+)", re.IGNORECASE),
    "perm_reg_no": re.compile(r"PERM\.?\s*REG\.?\s*NO\.?\s*[:\-]?\s*([A-Z0-9\/]+)", re.IGNORECASE),
    "student_name": re.compile(r"\bNAME\b\s*[:\-]?\s*([A-Z][A-Z.'\- ]{2,60})", re.IGNORECASE),
    "mother_name": re.compile(r"\bMOTHER\b\s*[:\-]?\s*([A-Z][A-Z.'\- ]{2,60})", re.IGNORECASE),
    "college": re.compile(r"\bCOLLEGE\b\s*[:\-]?\s*([A-Z][A-Z.,'\- ]{2,80})", re.IGNORECASE),
    "grand_total": re.compile(r"GRAND\s*.?TOTAL\s*=?\s*(\d+\s*/\s*\d+)", re.IGNORECASE),
    "result": re.compile(r"\bRESULT\b\s*[:\-]?\s*([A-Z][A-Z ]{2,40})", re.IGNORECASE),
    "invoice_number": re.compile(r"Invoice\s*(?:No\.?|Number|#)\s*[:\-]?\s*([A-Za-z0-9\-\/]+)", re.IGNORECASE),
    "amount_due": re.compile(r"Amount\s*Due\s*[:\-]?\s*([₹$€£]\s?[\d,]+\.?\d{0,2})", re.IGNORECASE),
    "bill_to": re.compile(r"Bill\s*To\s*[:\-]?\s*([^\n]+)", re.IGNORECASE),
}
_PHONE_CANDIDATE_RE = re.compile(
    r"(?:\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{3,4}(?:[-.\s]?\d{0,4})?"
)
FIELD_PRIORITY = [
    "invoice_number", "bill_to", "amount_due",
    "emails", "urls", "gstin_numbers", "pan_numbers", "ifsc_codes",
    "passport_numbers", "seat_no", "perm_reg_no", "grand_total",
    "student_name", "mother_name", "college", "result",
    "vehicle_reg_numbers", "aadhaar_numbers",
    "dates", "amounts", "pincodes", "phones",
]
FIELD_CONFIDENCE = {
    "emails": "high", "urls": "high", "gstin_numbers": "high",
    "pan_numbers": "high", "ifsc_codes": "high", "passport_numbers": "high",
    "seat_no": "high", "perm_reg_no": "high", "grand_total": "high",
    "student_name": "medium", "mother_name": "medium", "college": "medium",
    "result": "medium",
    "invoice_number": "high", "amount_due": "high", "bill_to": "medium",
    "vehicle_reg_numbers": "medium", "aadhaar_numbers": "medium",
    "dates": "medium", "amounts": "medium", "pincodes": "medium", "phones": "medium",
}
TAXONOMY = {
    "Personal": {
        "Identity & Legal": ["Aadhaar Card", "PAN Card", "Passport", "Driving Licence",
                             "Voter ID", "Birth Certificate", "Marriage Certificate",
                             "Name Change Affidavit", "Ration Card"],
        "Government Certificates": ["Domicile Certificate", "Caste Certificate",
                                   "Income Certificate", "Non-Creamy Layer Certificate"],
        "Education": ["10th Marksheet", "12th Marksheet", "Semester Marksheet",
                       "Hall Ticket / Admit Card", "Transfer Certificate (TC)",
                       "Bonafide Certificate", "School/College ID Card"],
        "Utilities & Bills": ["Electricity Bill", "Water Bill", "Gas Bill",
                              "Broadband/Internet Bill", "Mobile Bill"],
        "Financial": ["Bank Account Documents", "Fixed Deposits (FDs)", "Mutual Funds (MF)",
                       "IT Returns", "Form 16", "Loan Documents"],
        "Insurance": ["Health Insurance", "Life Insurance", "Vehicle Insurance",
                       "Home Insurance", "Travel Insurance"],
        "Healthcare": ["Medical Reports", "Prescriptions", "Vaccinations",
                        "Blood Group Information"],
        "Property": ["Purchase Documents", "Sale Deed", "Lease Agreement", "Property Tax"],
        "Vehicle": ["RC Book", "PUC", "Service History", "Purchase Warranty", "Road Tax"],
        "Gadgets/Appliances": ["Refrigerator", "Washing Machine", "TV", "Laptop",
                                "Robo Cleaner", "Purchase Invoice", "Warranty",
                                "User Manual", "AMC", "Service Record"],
        "Jewellery": ["Purchase Invoice", "Insurance", "Hallmark Certificate", "Warranty",
                      "Valuation Certificate"],
        "Travel": ["Flight Ticket", "Hotel Booking", "Visa", "Travel Insurance",
                   "Foreign Exchange Records"],
        "Other": ["Other"],
    },
    "Professional": {
        "Employment": ["Appointment Letter / Offer Letter", "Experience Certificate",
                       "Relieving Letter", "Salary Slip", "Promotion Letters", "Appraisal"],
        "Certification": ["AI Course", "Degree", "Memberships", "Sports", "Music"],
        "IP (Intellectual Property)": ["Patent Application", "Granted Patent",
                                       "Trademark", "Copyright"],
        "Business": ["GST Documents", "Company Registration", "MSME", "PAN", "TAN", "Licenses"],
        "Awards & Recognition": ["Awards", "Certificates", "Recognition Documents"],
        "Other": ["Other"],
    },
    "Corporate": {
        "Company Formation & Registration": ["Certificate of Incorporation", "MOA", "AOA",
                                             "Registration Documents"],
        "Board & Shareholder Documents": ["Board Resolutions", "Shareholder Resolutions",
                                          "Meeting Minutes", "Share Certificates"],
        "Corporate Governance & Compliance": ["Compliance Documents", "Corporate Policies",
                                               "Annual Filings", "Statutory Registers"],
        "Contracts & Commercial": ["Client Contracts", "Vendor Agreements",
                                   "Service Agreements", "Purchase Agreements"],
        "Finance & Tax": ["GST Documents", "Tax Documents", "Audited Financial Statements",
                           "Bank Documents"],
        "Intellectual Property": ["Trademark Documents", "Copyright Documents",
                                  "Patent Documents", "IP Assignment Agreements"],
        "Other": ["Other"],
    },
    "Legal": {
        "Client Management": ["Client Profiles", "Identity Proof", "Engagement Letters"],
        "Court Documents": ["Evidence", "Orders", "Affidavits", "Petitions"],
        "Agreements": ["Sale", "Lease", "Employment", "NDA", "Partnership",
                        "Proprietorship", "MOUs"],
        "Patents": ["Patent Documents"],
        "Corporate Legal": ["Corporate Legal Documents"],
        "Litigation Calendar": ["Litigation-related Calendar Documents"],
        "Other": ["Other"],
    },
}
MANUAL_KEYWORDS = {
    "Aadhaar Card": ["aadhaar", "aadhar", "uidai", "unique identification authority"],
    "PAN Card": ["income tax department", "permanent account number", "pan card"],
    "Passport": ["republic of india", "passport no", "given name", "place of birth"],
    "Driving Licence": ["driving licence", "driving license", "transport department", "dl no"],
    "Voter ID": ["election commission of india", "voter id", "epic no",
                 "elector's photo identity card"],
    "Birth Certificate": ["birth certificate", "registrar of births"],
    "Marriage Certificate": ["marriage certificate", "solemnized", "marriage registration"],
    "Name Change Affidavit": ["affidavit", "name change", "formerly known as"],
    "Ration Card": ["ration card", "public distribution", "fair price shop"],
    "Domicile Certificate": ["domicile certificate", "permanent resident"],
    "Caste Certificate": ["caste certificate", "scheduled caste", "scheduled tribe", "obc"],
    "Income Certificate": ["income certificate", "annual income", "tehsildar"],
    "Non-Creamy Layer Certificate": ["non-creamy layer", "non creamy layer", "ncl certificate"],
    "10th Marksheet": ["secondary school", "ssc", "matriculation", "10th"],
    "12th Marksheet": ["higher secondary", "hsc", "intermediate", "12th"],
    "Semester Marksheet": ["semester", "cgpa", "sgpa", "grade card",
                           "grand total", "seat no", "perm reg", "statement of marks",
                           "result", "first class", "second class", "distinction"],
    "Hall Ticket / Admit Card": ["admit card", "hall ticket", "roll no"],
    "Transfer Certificate (TC)": ["transfer certificate", "school leaving"],
    "Bonafide Certificate": ["bonafide certificate", "bona fide", "this is to certify"],
    "School/College ID Card": ["student id", "identity card", "class"],
    "Electricity Bill": ["electricity bill", "kwh", "consumer no", "discom", "units consumed"],
    "Water Bill": ["water bill", "water charges", "municipal corporation"],
    "Gas Bill": ["gas bill", "lpg", "cylinder", "gas agency"],
    "Broadband/Internet Bill": ["broadband", "internet bill", "data usage", "isp"],
    "Mobile Bill": ["mobile bill", "recharge", "call details", "telecom"],
    "Bank Account Documents": ["bank statement", "account number", "ifsc", "savings account"],
    "Fixed Deposits (FDs)": ["fixed deposit", "fd receipt", "maturity date"],
    "Mutual Funds (MF)": ["mutual fund", "folio number", "nav", "sip"],
    "IT Returns": ["income tax return", "itr", "acknowledgement number", "assessment year"],
    "Form 16": ["form 16", "tds", "salary certificate"],
    "Loan Documents": ["loan agreement", "emi", "sanction letter"],
    "Health Insurance": ["health insurance", "sum insured", "mediclaim"],
    "Life Insurance": ["life insurance", "policy holder", "sum assured"],
    "Vehicle Insurance": ["vehicle insurance", "motor insurance", "idv"],
    "Home Insurance": ["home insurance", "property insurance", "dwelling"],
    "Travel Insurance": ["travel insurance", "trip cover"],
    "Medical Reports": ["medical report", "diagnosis", "lab report", "pathology", "radiology"],
    "Prescriptions": ["prescription", "dosage", "tablet", "dr."],
    "Vaccinations": ["vaccination", "vaccine", "dose", "immunization", "cowin"],
    "Blood Group Information": ["blood group", "blood type"],
    "Sale Deed": ["sale deed", "vendor", "vendee", "registered deed"],
    "Lease Agreement": ["lease agreement", "lessor", "lessee", "rent agreement", "tenancy"],
    "Property Tax": ["property tax", "municipal tax"],
    "RC Book": ["registration certificate", "rc book", "chassis number", "engine number"],
    "PUC": ["pollution under control", "puc", "emission test"],
    "Service History": ["service history", "service record", "odometer", "workshop"],
    "GST Documents": ["gstin", "goods and services tax", "gst return", "gst certificate"],
    "Company Registration": ["certificate of incorporation", "registrar of companies", "cin"],
    "MSME": ["udyam", "msme registration", "udyog aadhaar"],
    "TAN": ["tax deduction account number", "tan"],
    "Certificate of Incorporation": ["certificate of incorporation", "cin",
                                     "registrar of companies"],
    "MOA": ["memorandum of association"],
    "AOA": ["articles of association"],
    "Board Resolutions": ["board resolution", "resolved that", "board of directors"],
    "Shareholder Resolutions": ["shareholder resolution", "special resolution",
                                 "ordinary resolution"],
    "Meeting Minutes": ["minutes of meeting", "agenda", "attendees"],
    "Share Certificates": ["share certificate", "folio no", "distinctive numbers"],
    "Compliance Documents": ["compliance certificate", "statutory compliance"],
    "Annual Filings": ["annual return", "roc filing"],
    "Statutory Registers": ["statutory register", "register of members"],
    "Client Contracts": ["contract agreement", "parties agree"],
    "Vendor Agreements": ["vendor agreement", "supplier agreement"],
    "Service Agreements": ["service agreement", "scope of services"],
    "Purchase Agreements": ["purchase order", "purchase agreement"],
    "Audited Financial Statements": ["balance sheet", "profit and loss", "auditor's report"],
    "Bank Documents": ["bank statement", "account number", "ifsc"],
    "Trademark Documents": ["trademark", "trade mark registry", "tm application"],
    "Copyright Documents": ["copyright", "copyright office"],
    "Patent Documents": ["patent", "patent application", "claims"],
    "Granted Patent": ["granted patent", "letters patent"],
    "Patent Application": ["patent application", "provisional specification"],
    "IP Assignment Agreements": ["assignment agreement", "assignor", "assignee"],
    "Appointment Letter / Offer Letter": ["offer letter", "appointment letter",
                                           "date of joining"],
    "Experience Certificate": ["experience certificate", "period of employment"],
    "Relieving Letter": ["relieving letter", "last working day"],
    "Salary Slip": ["salary slip", "payslip", "net pay", "gross salary"],
    "Promotion Letters": ["promotion letter", "promoted to"],
    "Appraisal": ["appraisal", "performance review"],
    "Client Profiles": ["client profile", "kyc"],
    "Identity Proof": ["identity proof", "id proof"],
    "Engagement Letters": ["engagement letter", "terms of engagement"],
    "Evidence": ["exhibit no", "evidence"],
    "Orders": ["in the court of", "honourable", "order dated"],
    "Affidavits": ["affidavit", "solemnly affirm", "deponent"],
    "Petitions": ["petition", "petitioner", "respondent"],
    "NDA": ["non-disclosure agreement", "confidentiality agreement", "nda"],
    "MOUs": ["memorandum of understanding", "mou"],
    "Flight Ticket": ["boarding pass", "pnr", "departure", "airline"],
    "Hotel Booking": ["hotel booking", "check-in", "check-out", "reservation confirmation"],
    "Visa": ["visa", "consulate", "embassy"],
    "Foreign Exchange Records": ["foreign exchange", "forex", "currency exchange"],
}
_STOPWORDS = {"of", "the", "and", "or", "for", "to", "in", "a", "an"}
def _keywords_for(doc_type):
    if doc_type in MANUAL_KEYWORDS:
        return MANUAL_KEYWORDS[doc_type]
    tokens = re.findall(r"[a-zA-Z]+", doc_type.lower())
    tokens = [t for t in tokens if t not in _STOPWORDS and len(t) > 2]
    return [doc_type.lower()] + tokens
DOC_TYPE_TO_PATH = {
    doc_type: (domain, category)
    for domain, categories in TAXONOMY.items()
    for category, doc_types in categories.items()
    for doc_type in doc_types
    if doc_type != "Other"
}
ALL_DOC_TYPES = sorted(DOC_TYPE_TO_PATH.keys())
REQUIRED_FIELDS = {
    'Identity & Legal|Aadhaar Card': ['Full Name', 'Aadhaar Number', 'Date of Birth', 'Gender', 'Address', 'VID'],
    'Identity & Legal|PAN Card': ['Full Name', "Father's/Mother's Name", 'PAN Number', 'Date of Birth', 'Date of Issue'],
    'Identity & Legal|Passport': ['Full Name', 'Passport Number', 'Date of Birth', 'Gender', 'Nationality', 'Place of Birth', 'Place of Issue', 'Date of Issue', 'Date of Expiry', 'Father/Mother/Spouse Name'],
    'Identity & Legal|Driving Licence': ['Full Name', 'Driving Licence Number', 'Date of Birth', 'Gender', 'Address', 'Issue Date', 'Validity / Expiry Date', 'Vehicle Classes', 'Issuing Authority / RTO'],
    'Identity & Legal|Voter ID': ['Full Name', 'EPIC / Voter ID Number', 'Date of Birth / Age', 'Gender', 'Father/Mother/Spouse Name', 'Address', 'Assembly Constituency', 'Polling Station'],
    'Identity & Legal|Birth Certificate': ["Child's Full Name", 'Date of Birth', 'Time of Birth', 'Place of Birth', 'Gender', "Father's Name", "Mother's Name", 'Registration Number', 'Date of Registration', 'Issuing Authority'],
    'Identity & Legal|Marriage Certificate': ["Husband's Full Name", "Wife's Full Name", 'Date of Marriage', 'Place of Marriage', 'Marriage Registration Number', 'Registration Date', "Father's/Mother's Names", 'Issuing Authority'],
    'Identity & Legal|Name Change Affidavit': ['Old Name', 'New Name', 'Date of Birth', "Father/Husband's Name", 'Address', 'Reason for Name Change', 'Affidavit Number', 'Date', 'Notary / Authority Name'],
    'Identity & Legal|Ration Card': ['Head of Family Name', 'Ration Card Number', 'Family Members', 'Address', 'Card Type', 'Issuing Authority', 'Date of Issue'],
    'Identity & Legal|Other': ['Document Title', 'Document Number', 'Full Name', 'Date', 'Issuing Authority', 'Important Reference Number'],
    'Government Certificates|Domicile Certificate': ['Applicant Name', "Father's/Husband's Name", 'Address', 'State / District', 'Certificate Number', 'Issuing Authority', 'Date of Issue', 'Purpose'],
    'Government Certificates|Caste Certificate': ['Applicant Name', "Father's Name", 'Caste', 'Caste Category', 'Address', 'Certificate Number', 'Issuing Authority', 'Date of Issue'],
    'Government Certificates|Income Certificate': ['Applicant Name', "Father's/Husband's Name", 'Annual Income', 'Financial Year', 'Address', 'Certificate Number', 'Issuing Authority', 'Date of Issue', 'Validity'],
    'Government Certificates|Non-Creamy Layer Certificate': ['Applicant Name', "Father's Name", 'Caste Category', 'Annual Income', 'Certificate Number', 'Issuing Authority', 'Date of Issue', 'Validity'],
    'Government Certificates|Other': ['Applicant Name', 'Certificate Title', 'Certificate Number', 'Issuing Authority', 'Date of Issue'],
    'Education|10th Marksheet': ['Student Full Name', 'School/Board Name', 'Class / Standard', 'Seat Number / Roll Number', 'Total Marks / Percentage', 'Grade', 'Subjects & Marks', 'Date of Issue', 'Year of Passing'],
    'Education|12th Marksheet': ['Student Full Name', 'School/Board Name', 'Class / Standard', 'Seat Number / Roll Number', 'Total Marks / Percentage', 'Grade', 'Subjects & Marks', 'Date of Issue', 'Year of Passing'],
    'Education|Semester Marksheet': ['Student Full Name', 'University / Institute Name', 'Course / Branch', 'Semester', 'Seat Number / Roll Number', 'PRN / Enrollment Number', 'SGPA / CGPA', 'Result', 'Date of Issue'],
    'Education|Hall Ticket / Admit Card': ['Student Full Name', "Father's/Mother's Name", 'Exam Name', 'Roll Number / Seat Number', 'Exam Centre', 'Exam Date', 'School/College Name', 'Date of Birth'],
    'Education|Transfer Certificate (TC)': ['Student Full Name', "Father's/Mother's Name", 'School/College Name', 'Date of Birth', 'Admission Number', 'TC Number', 'Date of Leaving', 'Class Last Studied', 'Reason for Leaving', 'Date of Issue'],
    'Education|Bonafide Certificate': ['Student Full Name', "Father's/Mother's Name", 'School/College Name', 'Class / Course', 'Academic Year', 'Purpose', 'Certificate Number', 'Date of Issue'],
    'Education|School/College ID Card': ['Student Full Name', 'School/College Name', 'Class / Course', 'Roll Number / ID Number', 'Blood Group', 'Valid Until', 'Contact Number'],
    'Education|Other': ['Student Full Name', 'School/College/University Name', 'Document Title', 'Reference Number', 'Date'],
    'Utilities & Bills|Electricity Bill': ['Consumer Name', 'Consumer Number', 'Bill Number', 'Billing Period', 'Units Consumed', 'Bill Amount', 'Due Date', 'Meter Number', 'Provider Name'],
    'Utilities & Bills|Water Bill': ['Consumer Name', 'Consumer Number', 'Bill Number', 'Billing Period', 'Bill Amount', 'Due Date', 'Provider Name'],
    'Utilities & Bills|Gas Bill': ['Consumer Name', 'Consumer Number', 'Bill Number', 'Billing Period', 'Bill Amount', 'Due Date', 'Provider Name'],
    'Utilities & Bills|Broadband/Internet Bill': ['Customer Name', 'Account Number', 'Bill Number', 'Billing Period', 'Plan Name', 'Bill Amount', 'Due Date', 'Provider Name'],
    'Utilities & Bills|Mobile Bill': ['Customer Name', 'Mobile Number', 'Bill Number', 'Billing Period', 'Plan Name', 'Bill Amount', 'Due Date', 'Provider Name'],
    'Utilities & Bills|Other': ['Customer Name', 'Account/Consumer Number', 'Bill Number', 'Bill Amount', 'Due Date', 'Provider Name'],
    'Financial|Bank Account Documents': ['Account Holder Name', 'Bank Name', 'Account Number', 'IFSC Code', 'Branch Name', 'Account Type', 'Customer ID', 'Statement Period', 'Balance'],
    'Financial|Fixed Deposits (FDs)': ['Account Holder Name', 'Bank Name', 'FD / Deposit Number', 'Principal Amount', 'Interest Rate', 'Start Date', 'Maturity Date', 'Maturity Amount', 'Tenure', 'Nominee Name'],
    'Financial|Mutual Funds (MF)': ['Investor Name', 'Folio Number', 'AMC / Fund House', 'Scheme Name', 'Plan / Option', 'Investment Amount', 'Units', 'NAV', 'Transaction Date', 'Current Value'],
    'Financial|IT Returns': ['Name', 'PAN', 'Assessment Year', 'Financial Year', 'ITR Form Type', 'Acknowledgement Number', 'Gross Total Income', 'Total Tax', 'Refund / Tax Payable', 'Filing Date'],
    'Financial|Form 16': ['Employee Name', 'PAN', 'Employer Name', 'Employer TAN', 'Assessment Year', 'Financial Year', 'Gross Salary', 'Tax Deducted (TDS)', 'Taxable Income', 'Date of Issue'],
    'Financial|Loan Documents': ['Borrower Name', 'Lender / Bank Name', 'Loan Account Number', 'Loan Type', 'Loan Amount', 'Interest Rate', 'Tenure', 'EMI Amount', 'Loan Start Date', 'Maturity / End Date', 'Outstanding Amount'],
    'Financial|Other': ['Account Holder / Customer Name', 'Institution Name', 'Account / Reference Number', 'Amount', 'Transaction / Issue Date', 'Maturity / Due Date', 'Important Reference Number'],
    'Insurance|Health Insurance': ['Policyholder Name', 'Policy Number', 'Insurance Company', 'Insured Person(s)', 'Policy Type', 'Sum Insured', 'Premium', 'Policy Start Date', 'Policy Expiry Date', 'TPA / Network Details'],
    'Insurance|Life Insurance': ['Policyholder Name', 'Life Assured Name', 'Policy Number', 'Insurance Company', 'Policy Type', 'Sum Assured', 'Premium', 'Policy Start Date', 'Maturity Date', 'Nominee Name'],
    'Insurance|Vehicle Insurance': ['Owner Name', 'Vehicle Registration Number', 'Policy Number', 'Insurance Company', 'Policy Type', 'Policy Start Date', 'Policy Expiry Date', 'IDV', 'Premium', 'Vehicle Make/Model'],
    'Insurance|Home Insurance': ['Policyholder Name', 'Policy Number', 'Insurance Company', 'Property Address', 'Policy Type', 'Sum Insured', 'Premium', 'Start Date', 'Expiry Date'],
    'Insurance|Travel Insurance': ['Traveller Name', 'Policy Number', 'Insurance Company', 'Destination', 'Trip Start Date', 'Trip End Date', 'Coverage Amount', 'Premium', 'Policy Start Date', 'Policy Expiry Date'],
    'Insurance|Other': ['Policyholder Name', 'Policy Number', 'Insurance Company', 'Policy Type', 'Sum Insured', 'Premium', 'Start Date', 'Expiry Date'],
    'Healthcare|Medical Reports': ['Patient Name', 'Patient ID', 'Date of Birth / Age', 'Gender', 'Doctor Name', 'Hospital / Lab Name', 'Report Type', 'Test Date', 'Report Date', 'Test Results', 'Diagnosis / Impression'],
    'Healthcare|Prescriptions': ['Patient Name', 'Doctor Name', 'Hospital / Clinic', 'Prescription Date', 'Medicine Names', 'Dosage', 'Frequency', 'Duration', 'Instructions'],
    'Healthcare|Vaccinations': ['Patient Name', 'Date of Birth', 'Vaccine Name', 'Dose Number', 'Vaccination Date', 'Next Dose Date', 'Batch Number', 'Vaccination Centre', 'Doctor / Healthcare Provider'],
    'Healthcare|Blood Group Information': ['Patient Name', 'Blood Group', 'Rh Factor', 'Test Date', 'Patient ID', 'Hospital / Lab Name'],
    'Healthcare|Other': ['Patient Name', 'Healthcare Provider', 'Document Type', 'Date', 'Diagnosis / Result', 'Important Medical Reference Number'],
    'Property|Purchase Documents': ['Buyer Name', 'Seller Name', 'Property Address', 'Property Type', 'Purchase Date', 'Purchase Amount', 'Registration Number', 'Survey / Plot Number', 'Document Number', 'Registration Authority'],
    'Property|Sale Deed': ['Buyer Name', 'Seller Name', 'Property Address', 'Property Description', 'Sale Amount', 'Sale Date', 'Survey / Plot Number', 'Registration Number', 'Registration Date', 'Registrar Office'],
    'Property|Lease Agreement': ['Landlord Name', 'Tenant Name', 'Property Address', 'Lease Start Date', 'Lease End Date', 'Monthly Rent', 'Security Deposit', 'Lease Duration', 'Agreement Number', 'Notice Period'],
    'Property|Property Tax': ['Property Owner Name', 'Property Address', 'Property ID / Assessment Number', 'Tax Year', 'Tax Amount', 'Paid Amount', 'Due Date', 'Payment Date', 'Receipt Number'],
    'Property|Other': ['Owner Name', 'Property Address', 'Property ID / Survey Number', 'Document Number', 'Transaction Amount', 'Date', 'Registration Authority'],
    'Vehicle|RC Book': ['Owner Name', 'Registration Number', 'Vehicle Make', 'Vehicle Model', 'Vehicle Class', 'Fuel Type', 'Chassis Number', 'Engine Number', 'Registration Date', 'Manufacturing Date', 'RC Validity', 'RTO / Registering Authority'],
    'Vehicle|PUC': ['Vehicle Registration Number', 'Owner Name', 'PUC Certificate Number', 'Fuel Type', 'Emission Test Date', 'Validity / Expiry Date', 'Emission Values', 'Testing Centre'],
    'Vehicle|Service History': ['Owner Name', 'Vehicle Registration Number', 'Vehicle Make / Model', 'Service Centre', 'Service Date', 'Odometer Reading', 'Service Type', 'Parts Replaced', 'Service Cost', 'Next Service Due'],
    'Vehicle|Purchase Warranty': ['Owner / Customer Name', 'Vehicle Make / Model', 'VIN / Chassis Number', 'Purchase Date', 'Warranty Start Date', 'Warranty End Date', 'Dealer Name', 'Warranty Number'],
    'Vehicle|Road Tax': ['Owner Name', 'Vehicle Registration Number', 'Tax Receipt Number', 'Tax Amount', 'Tax Period', 'Payment Date', 'Validity / Expiry Date', 'RTO'],
    'Vehicle|Other': ['Owner Name', 'Registration Number', 'Vehicle Make / Model', 'Document Number', 'Date', 'Expiry Date', 'Important Reference Number'],
    'Gadgets/Appliances|Refrigerator': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Product ID', 'Purchase Date', 'Dealer / Seller', 'Invoice Number', 'Warranty Start Date', 'Warranty End Date'],
    'Gadgets/Appliances|Washing Machine': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Product ID', 'Purchase Date', 'Dealer / Seller', 'Invoice Number', 'Warranty Start Date', 'Warranty End Date'],
    'Gadgets/Appliances|TV': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Product ID', 'Purchase Date', 'Dealer / Seller', 'Invoice Number', 'Warranty Start Date', 'Warranty End Date'],
    'Gadgets/Appliances|Laptop': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Product ID', 'Purchase Date', 'Dealer / Seller', 'Invoice Number', 'Warranty Start Date', 'Warranty End Date'],
    'Gadgets/Appliances|Robo Cleaner': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Product ID', 'Purchase Date', 'Dealer / Seller', 'Invoice Number', 'Warranty Start Date', 'Warranty End Date'],
    'Gadgets/Appliances|Purchase Invoice': ['Customer Name', 'Seller Name', 'Invoice Number', 'Invoice Date', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Quantity', 'Purchase Amount', 'Tax Amount', 'Total Amount'],
    'Gadgets/Appliances|Warranty': ['Customer Name', 'Brand', 'Product Name', 'Model Number', 'Serial Number', 'Warranty Number', 'Warranty Start Date', 'Warranty End Date', 'Dealer / Service Provider'],
    'Gadgets/Appliances|User Manual': ['Product Name', 'Brand', 'Model Number', 'Product Type', 'Manual Version / Edition', 'Important Product Reference'],
    'Gadgets/Appliances|AMC': ['Customer Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'AMC Number', 'AMC Start Date', 'AMC End Date', 'Service Provider', 'Amount'],
    'Gadgets/Appliances|Service Record': ['Customer Name', 'Product Name', 'Model Number', 'Serial Number', 'Service Centre', 'Service Date', 'Issue / Complaint', 'Service Performed', 'Parts Replaced', 'Service Cost', 'Next Service Date'],
    'Gadgets/Appliances|Other': ['Customer / Owner Name', 'Product Name', 'Brand', 'Model Number', 'Serial Number', 'Purchase Date', 'Warranty Expiry', 'Invoice / Reference Number'],
    'Jewellery|Purchase Invoice': ['Customer Name', 'Jeweller Name', 'Invoice Number', 'Invoice Date', 'Jewellery Description', 'Metal Type', 'Purity', 'Gross Weight', 'Net Weight', 'Stone Weight', 'Making Charges', 'Tax', 'Total Amount'],
    'Jewellery|Insurance': ['Policyholder Name', 'Policy Number', 'Jewellery Description', 'Insured Value', 'Insurance Company', 'Start Date', 'Expiry Date', 'Premium'],
    'Jewellery|Hallmark Certificate': ['Owner / Customer Name', 'Jewellery Description', 'Hallmark / HUID Number', 'Metal Type', 'Purity', 'Weight', 'Jeweller Name', 'Hallmark Centre', 'Certificate Date'],
    'Jewellery|Warranty': ['Customer Name', 'Jewellery Description', 'Jeweller Name', 'Invoice Number', 'Warranty Number', 'Purchase Date', 'Warranty Start Date', 'Warranty End Date'],
    'Jewellery|Valuation Certificate': ['Owner Name', 'Jewellery Description', 'Metal Type', 'Purity', 'Gross Weight', 'Net Weight', 'Stone Details', 'Valuation Amount', 'Valuation Date', 'Valuer Name', 'Certificate Number'],
    'Jewellery|Other': ['Owner Name', 'Jewellery Description', 'Metal Type', 'Purity', 'Weight', 'Value', 'Date', 'Certificate / Invoice Number'],
    'Travel|Flight Ticket': ['Passenger Name', 'PNR', 'Ticket Number', 'Airline', 'Flight Number', 'Departure Airport', 'Arrival Airport', 'Departure Date', 'Departure Time', 'Arrival Date', 'Arrival Time', 'Seat Number', 'Class', 'Booking Reference'],
    'Travel|Hotel Booking': ['Guest Name', 'Hotel Name', 'Booking Number', 'Check-in Date', 'Check-out Date', 'Number of Guests', 'Room Type', 'Number of Rooms', 'Booking Amount', 'Address', 'Contact Number'],
    'Travel|Visa': ['Applicant Name', 'Passport Number', 'Visa Number', 'Nationality', 'Date of Birth', 'Visa Type', 'Country', 'Issue Date', 'Expiry Date', 'Number of Entries', 'Duration of Stay'],
    'Travel|Travel Insurance': ['Traveller Name', 'Policy Number', 'Insurance Company', 'Destination', 'Trip Start Date', 'Trip End Date', 'Coverage Amount', 'Premium', 'Policy Start Date', 'Policy Expiry Date'],
    'Travel|Foreign Exchange Records': ['Customer Name', 'Transaction Number', 'Exchange Provider / Bank', 'Currency', 'Foreign Currency Amount', 'Exchange Rate', 'INR Amount', 'Transaction Date', 'Receipt Number'],
    'Travel|Other': ['Traveller Name', 'Booking / Reference Number', 'Destination', 'Travel Date', 'Return Date', 'Provider', 'Amount', 'Important Reference Number'],
    'Employment|Appointment Letter / Offer Letter': ['Employee Name', 'Employer / Company Name', 'Designation', 'Department', 'Date of Joining', 'Offer Date', 'CTC / Salary', 'Employment Type', 'Reporting Manager', 'Reference Number'],
    'Employment|Experience Certificate': ['Employee Name', 'Employer / Company Name', 'Designation', 'Department', 'Date of Joining', 'Date of Relieving', 'Duration of Employment', 'Certificate Date', 'Reference Number'],
    'Employment|Relieving Letter': ['Employee Name', 'Employer / Company Name', 'Designation', 'Date of Joining', 'Last Working Day', 'Relieving Date', 'Reference Number'],
    'Employment|Salary Slip': ['Employee Name', 'Employer / Company Name', 'Employee ID', 'Designation', 'Pay Period', 'Gross Salary', 'Deductions', 'Net Salary', 'Date of Issue'],
    'Employment|Promotion Letters': ['Employee Name', 'Employer / Company Name', 'Previous Designation', 'New Designation', 'Effective Date', 'New Salary / CTC', 'Reference Number'],
    'Employment|Appraisal': ['Employee Name', 'Employer / Company Name', 'Designation', 'Appraisal Period', 'Rating', 'Revised Salary / CTC', 'Effective Date'],
    'Employment|Other': ['Employee Name', 'Employer / Company Name', 'Designation', 'Date', 'Reference Number'],
    'Certification|Degree': ['Student Full Name', 'University / Institute Name', 'Degree Name', 'Branch / Specialization', 'Seat Number / Roll Number', 'Enrollment Number / PRN', 'Class / CGPA / Percentage', 'Certificate Number', 'Date of Issue', 'Year of Passing'],
    'Certification|AI Course': ['Student / Participant Name', 'Course Name', 'Institution / Platform Name', 'Certificate Number', 'Completion Date', 'Grade / Score', 'Duration'],
    'Certification|Memberships': ['Member Name', 'Organization / Association Name', 'Membership Number', 'Membership Type', 'Valid From', 'Valid Until'],
    'Certification|Sports': ['Participant Name', 'Event / Tournament Name', 'Organizing Authority', 'Position / Achievement', 'Date', 'Certificate Number'],
    'Certification|Music': ['Participant Name', 'Course / Exam Name', 'Institution Name', 'Grade / Level', 'Certificate Number', 'Date of Issue'],
    'Certification|Others': ['Recipient Name', 'Certificate Title', 'Issuing Authority', 'Certificate Number', 'Date of Issue'],
    'IP (Intellectual Property)|Patent Application': ['Applicant Name', 'Invention Title', 'Application Number', 'Filing Date', 'Patent Office', 'Status'],
    'IP (Intellectual Property)|Granted Patent': ['Patentee Name', 'Invention Title', 'Patent Number', 'Filing Date', 'Grant Date', 'Patent Office', 'Validity'],
    'IP (Intellectual Property)|Trademark': ['Applicant / Owner Name', 'Trademark Name', 'Application / Registration Number', 'Class', 'Filing Date', 'Registration Date', 'Validity'],
    'IP (Intellectual Property)|Copyright': ['Author / Owner Name', 'Work Title', 'Registration Number', 'Registration Date', 'Copyright Office'],
    'IP (Intellectual Property)|Other': ['Applicant Name', 'Reference Number', 'Filing Date', 'Issuing Authority'],
    'Business|GST Documents': ['Business Name', 'GSTIN', 'Registration Date', 'Business Address', 'Constitution of Business'],
    'Business|Company Registration': ['Company Name', 'CIN', 'Registration Date', 'Registered Address', 'Registrar of Companies'],
    'Business|MSME': ['Business Name', 'Udyam / MSME Registration Number', 'Registration Date', 'Business Category'],
    'Business|PAN': ['Entity Name', 'PAN Number', 'Date of Issue'],
    'Business|TAN': ['Entity Name', 'TAN Number', 'Date of Issue'],
    'Business|Licenses': ['Business Name', 'License Number', 'Issuing Authority', 'Issue Date', 'Expiry Date'],
    'Business|Other': ['Business Name', 'Reference Number', 'Issuing Authority', 'Date'],
    'Awards & Recognition|Awards': ['Recipient Name', 'Award Title', 'Awarding Organization', 'Date', 'Category / Field'],
    'Awards & Recognition|Certificates': ['Recipient Name', 'Certificate Title', 'Issuing Authority', 'Certificate Number', 'Date of Issue'],
    'Awards & Recognition|Recognition Documents': ['Recipient Name', 'Recognition Title', 'Issuing Organization', 'Date'],
    'Awards & Recognition|Other': ['Recipient Name', 'Reference Number', 'Issuing Authority', 'Date'],
}
_KEYWORD_PATTERNS = {
    doc_type: [re.compile(r"\b" + re.escape(kw) + r"\b") for kw in _keywords_for(doc_type)]
    for doc_type in ALL_DOC_TYPES
}
_OTHER_RESULT = {"domain": "Other", "category": "Other", "document_type": "Other",
                 "confidence": 0.0, "method": "keyword"}
_LOW_CONFIDENCE_FLOOR = 0.15
def _keyword_classify(raw_text):
    text_lower = raw_text.lower()
    best = None
    best_score = 0
    for doc_type, patterns in _KEYWORD_PATTERNS.items():
        score = sum(1 for pat in patterns if pat.search(text_lower))
        if score > best_score:
            best_score = score
            best = doc_type
    if best is None:
        return dict(_OTHER_RESULT)
    domain, category = DOC_TYPE_TO_PATH[best]
    normalized = round(min(best_score / 4.0, 1.0), 2)
    if normalized < _LOW_CONFIDENCE_FLOOR:
        return dict(_OTHER_RESULT)
    return {"domain": domain, "category": category, "document_type": best,
            "confidence": normalized, "method": "keyword"}
DEFAULT_MODEL_PATH = "classifier_model.joblib"
DEFAULT_FEEDBACK_PATH = "training_data.jsonl"
DEFAULT_LOW_TIER_DB_PATH = "low_tier_value_store.json"
LOW_TIER_MATCH_THRESHOLD = 0.6
def _synthesize_seed_examples(per_class=6, seed=42):
    rng = random.Random(seed)
    templates = [
        "{kws}",
        "This document is a {name}. {kws}",
        "{name}\n{kws}",
        "Issued document: {name}. Details: {kws}",
        "{kws}\n{name}",
    ]
    examples = []
    for doc_type in ALL_DOC_TYPES:
        kws = _keywords_for(doc_type)
        for _ in range(per_class):
            k = kws if len(kws) <= 3 else rng.sample(kws, k=rng.randint(2, min(4, len(kws))))
            text = rng.choice(templates).format(name=doc_type, kws=". ".join(k))
            examples.append({"text": text, "document_type": doc_type})
    return examples
def _load_jsonl(path):
    examples = []
    if not os.path.exists(path):
        return examples
    with open(path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                row = json.loads(line)
            except json.JSONDecodeError:
                continue
            if row.get("text") and row.get("document_type") in DOC_TYPE_TO_PATH:
                examples.append(row)
    return examples
def add_training_example(text, document_type, feedback_path=DEFAULT_FEEDBACK_PATH):
    if document_type not in DOC_TYPE_TO_PATH:
        raise ValueError(
            f"'{document_type}' is not a known document_type. "
            f"See ocr_pipeline.ALL_DOC_TYPES for valid values."
        )
    with open(feedback_path, "a", encoding="utf-8") as f:
        f.write(json.dumps({"text": text, "document_type": document_type},
                            ensure_ascii=False) + "\n")
def train_classifier(feedback_path=DEFAULT_FEEDBACK_PATH, model_path=DEFAULT_MODEL_PATH,
                     seed_per_class=6):
    if not _ML_AVAILABLE:
        raise RuntimeError(
            "scikit-learn/joblib not installed. Run: "
            "pip install scikit-learn joblib --break-system-packages"
        )
    seed_examples = _synthesize_seed_examples(per_class=seed_per_class)
    real_examples = _load_jsonl(feedback_path)
    weighted_examples = seed_examples + real_examples * 5
    texts = [ex["text"] for ex in weighted_examples]
    labels = [ex["document_type"] for ex in weighted_examples]
    pipeline = Pipeline([
        ("tfidf", TfidfVectorizer(ngram_range=(1, 2), min_df=1, sublinear_tf=True)),
        ("clf", LogisticRegression(max_iter=1000)),
    ])
    pipeline.fit(texts, labels)
    joblib.dump(pipeline, model_path)
    return {
        "model_path": model_path,
        "seed_examples": len(seed_examples),
        "real_examples": len(real_examples),
        "classes": len(set(labels)),
    }
def load_classifier(model_path=DEFAULT_MODEL_PATH):
    if not _ML_AVAILABLE or not os.path.exists(model_path):
        return None
    try:
        return joblib.load(model_path)
    except Exception:
        return None
DEFAULT_AUTO_LEARN_STATE_PATH = "auto_learn_state.json"
def _load_auto_learn_state(state_path=DEFAULT_AUTO_LEARN_STATE_PATH):
    if os.path.exists(state_path):
        try:
            with open(state_path, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            pass
    return {"total_auto_logged": 0, "examples_since_retrain": 0, "last_retrained_at": None}
def _save_auto_learn_state(state, state_path=DEFAULT_AUTO_LEARN_STATE_PATH):
    with open(state_path, "w", encoding="utf-8") as f:
        json.dump(state, f)
def auto_learn(raw_text, classification, feedback_path=DEFAULT_FEEDBACK_PATH,
               model_path=DEFAULT_MODEL_PATH, state_path=DEFAULT_AUTO_LEARN_STATE_PATH,
               min_confidence=0.75, retrain_every=25):
    if classification["method"] != "hybrid" or classification["confidence"] < min_confidence:
        return {"auto_logged": False, "retrained": False}
    if classification["document_type"] == "Other" or not raw_text.strip():
        return {"auto_logged": False, "retrained": False}
    add_training_example(raw_text, classification["document_type"], feedback_path=feedback_path)
    state = _load_auto_learn_state(state_path)
    state["total_auto_logged"] += 1
    state["examples_since_retrain"] += 1
    retrained = False
    if state["examples_since_retrain"] >= retrain_every:
        train_classifier(feedback_path=feedback_path, model_path=model_path)
        state["examples_since_retrain"] = 0
        state["last_retrained_at"] = datetime.now(timezone.utc).isoformat()
        retrained = True
    _save_auto_learn_state(state, state_path)
    return {"auto_logged": True, "retrained": retrained, "total_auto_logged": state["total_auto_logged"]}
def _ml_classify(raw_text, model):
    proba = model.predict_proba([raw_text])[0]
    classes = model.classes_
    top_idx = int(np.argmax(proba))
    doc_type = str(classes[top_idx])
    confidence = float(proba[top_idx])
    domain, category = DOC_TYPE_TO_PATH.get(doc_type, ("Other", "Other"))
    return {"domain": domain, "category": category, "document_type": doc_type,
            "confidence": round(confidence, 2), "method": "ml"}
def classify_document(raw_text, model=None, min_ml_confidence=0.35):
    keyword_result = _keyword_classify(raw_text)
    if model is None:
        model = load_classifier()
    if model is None or not raw_text.strip():
        return keyword_result
    ml_result = _ml_classify(raw_text, model)
    if ml_result["confidence"] < _LOW_CONFIDENCE_FLOOR:
        ml_result = dict(_OTHER_RESULT, method="ml")
    if ml_result["document_type"] == "Other" and keyword_result["document_type"] == "Other":
        return keyword_result
    if ml_result["document_type"] == keyword_result["document_type"]:
        boosted = min(1.0, round((ml_result["confidence"] + keyword_result["confidence"]) / 2
                                 + 0.15, 2))
        return {**ml_result, "confidence": boosted, "method": "hybrid"}
    if ml_result["confidence"] < min_ml_confidence and keyword_result["document_type"] != "Other":
        return keyword_result
    return ml_result if ml_result["confidence"] >= keyword_result["confidence"] else keyword_result
def extract_fields(raw_text):
    claimed_spans = []
    def _overlaps(span):
        return any(s < span[1] and span[0] < e for s, e in claimed_spans)
    fields = {}
    for field_name in FIELD_PRIORITY:
        values = []
        if field_name == "phones":
            for m in _PHONE_CANDIDATE_RE.finditer(raw_text):
                if _overlaps(m.span()):
                    continue
                digit_count = len(re.sub(r"\D", "", m.group()))
                if 10 <= digit_count <= 13:
                    claimed_spans.append(m.span())
                    values.append(m.group().strip())
        else:
            for m in FIELD_PATTERNS[field_name].finditer(raw_text):
                if _overlaps(m.span()):
                    continue
                claimed_spans.append(m.span())
                value = m.group(1) if m.groups() else m.group()
                values.append(value.strip())
        if values:
            fields[field_name] = {
                "values": list(dict.fromkeys(values)),
                "confidence": FIELD_CONFIDENCE.get(field_name, "medium"),
            }
    return fields
def _label_alternatives(label):
    parts = [p.strip() for p in re.split(r"\s*/\s*", label) if p.strip()]
    candidates = [label] + parts
    if len(parts) > 1:
        last_words = parts[-1].split()
        if len(last_words) > 1:
            suffix = last_words[-1]
            for p in parts[:-1]:
                if len(p.split()) == 1:
                    candidates.append(f"{p} {suffix}")
    expanded = []
    for phrase in candidates:
        expanded.append(phrase)
        stripped = re.sub(r"\(s\)", "", phrase).strip()
        if stripped and stripped != phrase:
            expanded.append(stripped)
        if re.search(r"\bNumber\b", phrase, re.IGNORECASE):
            expanded.append(re.sub(r"\bNumber\b", "No", phrase, flags=re.IGNORECASE))
    seen = set()
    result = []
    for phrase in expanded:
        key = phrase.lower()
        if key in seen or len(phrase) < 4:
            continue
        seen.add(key)
        result.append(phrase)
    result.sort(key=len, reverse=True)
    return result or [label]
def _build_label_regex(label):
    alt_patterns = []
    for alt in _label_alternatives(label):
        escaped = r"\s+".join(re.escape(word) for word in alt.split())
        alt_patterns.append(escaped)
    core = r"(?<!\w)(?:" + "|".join(alt_patterns) + r")(?!\w)"
    separator = r"(?:[ \t]*[:\-][ \t]*|[ \t]+)"
    pattern = core + separator + r"([^\n\t]{2,80}?)(?=\n|[ \t]{2,}|\t|$)"
    return re.compile(pattern, re.IGNORECASE)
_ALL_REQUIRED_LABELS = sorted({label for fields in REQUIRED_FIELDS.values() for label in fields})
_LABEL_PATTERNS = {label: _build_label_regex(label) for label in _ALL_REQUIRED_LABELS}
_ALL_REQUIRED_LABELS_NORM = {re.sub(r"[^a-z0-9]+", "", lbl.lower()) for lbl in _ALL_REQUIRED_LABELS}
def _is_probably_a_label(text):
    if not text:
        return False
    norm = re.sub(r"[^a-z0-9]+", "", text.lower())
    if not norm:
        return False
    if norm in _ALL_REQUIRED_LABELS_NORM:
        return True
    for lbl in _ALL_REQUIRED_LABELS:
        for part in re.split(r"\s*/\s*", lbl):
            if re.sub(r"[^a-z0-9]+", "", part.lower()) == norm:
                return True
    return False
def _build_label_only_regex(label):
    alt_patterns = []
    for alt in _label_alternatives(label):
        escaped = r"\s+".join(re.escape(word) for word in alt.split())
        alt_patterns.append(escaped)
    core = r"(?:" + "|".join(alt_patterns) + r")"
    pattern = r"^[ \t]*(?:" + core + r")[ \t]*[:\-]?[ \t]*$"
    return re.compile(pattern, re.IGNORECASE | re.MULTILINE)
_LABEL_ONLY_PATTERNS = {label: _build_label_only_regex(label) for label in _ALL_REQUIRED_LABELS}
def _next_nonempty_line(text, after_pos, max_lookahead_lines=3):
    rest = text[after_pos:]
    checked = 0
    for m in re.finditer(r"[ \t]*([^\n]*)", rest):
        line = m.group(1)
        if not line.strip():
            if m.end() >= len(rest):
                break
            continue
        checked += 1
        return line.strip(), (after_pos + m.start(1), after_pos + m.end(1))
    return None, None
def _match_label_next_line(text, label, claimed_spans, label_only_pattern=None):
    pattern = label_only_pattern or _build_label_only_regex(label)
    for m in pattern.finditer(text):
        label_span = m.span()
        if _span_overlaps(label_span, claimed_spans):
            continue
        value, value_span = _next_nonempty_line(text, m.end())
        if not value or _is_probably_a_label(value):
            continue
        if _span_overlaps(value_span, claimed_spans):
            continue
        claimed_spans.append(label_span)
        return value, value_span
    return None, None
def _span_overlaps(span, claimed_spans):
    return any(s < span[1] and span[0] < e for s, e in claimed_spans)
def _fallback_pattern(text, pattern, claimed_spans):
    for m in pattern.finditer(text):
        span = m.span()
        if _span_overlaps(span, claimed_spans):
            continue
        value = m.group(1) if m.groups() else m.group()
        value = value.strip()
        if value:
            return value, span
    return None, None
def _fallback_date(text, claimed_spans):
    return _fallback_pattern(text, FIELD_PATTERNS["dates"], claimed_spans)
def _fallback_gender(text, claimed_spans):
    pattern = re.compile(r"\b(Male|Female|Transgender)\b", re.IGNORECASE)
    return _fallback_pattern(text, pattern, claimed_spans)
_ADDRESS_HINT_RE = re.compile(
    r"\b(road|street|nagar|colony|sector|lane|marg|block|floor|apartment|"
    r"flat|house\s*no|village|tehsil|district|dist\.?|near|opp\.?|behind)\b",
    re.IGNORECASE,
)
def _fallback_address(text, claimed_spans):
    for m in re.finditer(r"^[ \t]*(.{8,120})[ \t]*$", text, re.MULTILINE):
        line = m.group(1).strip()
        if not line or len(line.split()) < 3:
            continue
        has_pin = bool(re.search(r"\b\d{6}\b", line))
        has_hint = bool(_ADDRESS_HINT_RE.search(line))
        if not (has_pin or has_hint):
            continue
        span = m.span(1)
        if _span_overlaps(span, claimed_spans):
            continue
        return line, span
    return None, None
_NAME_LINE_RE = re.compile(
    r"^[ \t]*([A-Z][A-Za-z.'\-]*(?:[ \t]+[A-Za-z.'\-]+){1,3})[ \t]*$", re.MULTILINE
)
_HONORIFIC_NAME_RE = re.compile(
    r"\b(?:Smt|Shri|Sri|Kumari|Miss|Mr|Mrs|Ms|Dr)\.?[ \t]+"
    r"([A-Z]{2,}(?:[ \t]+[A-Z]{2,}){1,3})\b"
)
_NAME_BLOCKLIST_WORDS = {
    "government", "india", "department", "income", "tax", "permanent",
    "account", "number", "card", "certificate", "republic", "authority",
    "ministry", "corporation", "university", "board", "school", "college",
    "signature", "address", "date", "issue", "valid", "validity", "male",
    "female", "gender", "birth", "identity", "proof", "original", "copy",
    "specimen", "sample", "form", "office", "district", "state", "branch",
    "bank", "policy", "invoice", "bill", "receipt", "report", "hospital",
    "advanced", "microprocessors", "engineering", "electronics",
    "communication", "network", "networks", "optical", "instrumentation",
    "measurements", "measurement", "systems", "system", "design", "power",
    "computer", "biomedical", "project", "seminar", "seminars", "term",
    "work", "part", "paper", "max", "min", "obt", "result", "class",
    "distinction", "institute", "examination", "statement", "showing",
    "candidate", "passing", "marks", "obtained", "semester", "branch",
    "video", "instrumentation",
}
def _fallback_name(text, claimed_spans):
    for m in _HONORIFIC_NAME_RE.finditer(text):
        span = m.span(1)
        if _span_overlaps(span, claimed_spans):
            continue
        return m.group(1).strip(), span
    for m in _NAME_LINE_RE.finditer(text):
        candidate = m.group(1).strip()
        words = candidate.split()
        if not (2 <= len(words) <= 4):
            continue
        if any(ch.isdigit() for ch in candidate):
            continue
        if any(w.strip(".'-").lower() in _NAME_BLOCKLIST_WORDS for w in words):
            continue
        span = m.span(1)
        if _span_overlaps(span, claimed_spans):
            continue
        return candidate, span
    return None, None
_INSTITUTION_RE = re.compile(
    r"\b([A-Z][A-Za-z&'.\- ]{1,40}(?:University|College|Institute|School|"
    r"Hospital|Bank|Corporation|Insurance|Limited))\b",
    re.IGNORECASE,
)
def _fallback_institution(text, claimed_spans):
    return _fallback_pattern(text, _INSTITUTION_RE, claimed_spans)
def _guess_field_type(label):
    lower = label.lower()
    if "aadhaar" in lower:
        return "aadhaar"
    if "gstin" in lower:
        return "gstin"
    if re.search(r"\bpan\b", lower):
        return "pan"
    if "passport" in lower and "number" in lower:
        return "passport"
    if "ifsc" in lower:
        return "ifsc"
    if "email" in lower:
        return "email"
    if any(kw in lower for kw in ("phone", "mobile", "contact number")):
        return "phone"
    if "gender" in lower or lower.strip() == "sex":
        return "gender"
    if "address" in lower:
        return "address"
    if "date" in lower or lower.strip() == "dob" or "age" in lower:
        return "date"
    if any(kw in lower for kw in ("amount", "premium", "salary", "income", "balance", "value")):
        return "amount"
    if any(kw in lower for kw in (
        "university", "institute", "college", "school", "hospital", "bank",
        "company", "employer", "provider", "insurance", "dealer", "seller",
        "corporation", "jeweller", "jeweler", "airline", "hotel",
        "organization", "organisation", "association", "authority",
        "office", "registrar", "business", "entity", "awarding",
        "issuing", "lab name", "clinic", "lender",
    )):
        return "institution"
    if "name" in lower:
        return "name"
    return None
_FALLBACK_EXTRACTORS = {
    "aadhaar": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["aadhaar_numbers"], claimed),
    "gstin": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["gstin_numbers"], claimed),
    "pan": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["pan_numbers"], claimed),
    "passport": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["passport_numbers"], claimed),
    "ifsc": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["ifsc_codes"], claimed),
    "email": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["emails"], claimed),
    "phone": lambda text, claimed: _fallback_pattern(text, _PHONE_CANDIDATE_RE, claimed),
    "amount": lambda text, claimed: _fallback_pattern(text, FIELD_PATTERNS["amounts"], claimed),
    "gender": _fallback_gender,
    "address": _fallback_address,
    "date": _fallback_date,
    "name": _fallback_name,
    "institution": _fallback_institution,
}
def load_value_db(db_path=DEFAULT_LOW_TIER_DB_PATH):
    if not os.path.exists(db_path):
        return {}
    try:
        with open(db_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, dict) else {}
    except (json.JSONDecodeError, OSError):
        return {}
def save_value_db(db, db_path=DEFAULT_LOW_TIER_DB_PATH):
    with open(db_path, "w", encoding="utf-8") as f:
        json.dump(db, f, indent=2, ensure_ascii=False, sort_keys=True)
def _fuzzy_ratio(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()
LOW_TIER_SAFE_FIELDS_TABLE = [
    ("Corporate", "Board & Shareholder Documents", "Board Resolutions", ["Company Name", "Resolution Title", "Subject Matter"]),
    ("Corporate", "Board & Shareholder Documents", "Meeting Minutes", ["Company Name", "Meeting Type", "Key Decisions"]),
    ("Corporate", "Board & Shareholder Documents", "Other", ["Company Name", "Document Title"]),
    ("Corporate", "Board & Shareholder Documents", "Share Certificates", ["Company Name", "Number of Shares", "Share Class/Type"]),
    ("Corporate", "Board & Shareholder Documents", "Shareholder Resolutions", ["Company Name", "Resolution Title", "Subject Matter"]),
    ("Corporate", "Company Formation & Registration", "AOA (Articles of Association)", ["Company Name", "Amendments Summary"]),
    ("Corporate", "Company Formation & Registration", "Certificate of Incorporation", ["Company Name", "Registrar of Companies (RoC)", "Company Type", "Authorized Capital"]),
    ("Corporate", "Company Formation & Registration", "MOA (Memorandum of Association)", ["Company Name", "Object Clause Summary", "Authorized Capital"]),
    ("Corporate", "Company Formation & Registration", "Other", ["Company Name", "Document Title", "Issuing Authority"]),
    ("Corporate", "Company Formation & Registration", "Registration Documents", ["Company/Entity Name", "Registering Authority"]),
    ("Corporate", "Contracts & Commercial", "Client Contracts", ["Company Name", "Contract Type", "Renewal Terms"]),
    ("Corporate", "Contracts & Commercial", "Other", ["Company Name"]),
    ("Corporate", "Contracts & Commercial", "Purchase Agreements", ["Company Name", "Goods/Services Description", "Delivery Terms"]),
    ("Corporate", "Contracts & Commercial", "Service Agreements", ["Company Name", "Service Description"]),
    ("Corporate", "Contracts & Commercial", "Vendor Agreements", ["Company Name", "Agreement Type", "Payment Terms"]),
    ("Corporate", "Corporate Governance & Compliance", "Annual Filings", ["Company Name", "Financial Year", "Filing Type"]),
    ("Corporate", "Corporate Governance & Compliance", "Compliance Documents", ["Company Name", "Compliance Type", "Regulatory Authority", "Filing Period"]),
    ("Corporate", "Corporate Governance & Compliance", "Corporate Policies", ["Company Name", "Policy Title", "Version"]),
    ("Corporate", "Corporate Governance & Compliance", "Other", ["Company Name", "Document Title"]),
    ("Corporate", "Corporate Governance & Compliance", "Statutory Registers", ["Company Name", "Register Type", "Maintained Since"]),
    ("Corporate", "Finance & Tax", "Audited Financial Statements", ["Company Name", "Financial Year"]),
    ("Corporate", "Finance & Tax", "Bank Documents", ["Company Name", "Bank Name", "Account Type", "Statement Period"]),
    ("Corporate", "Finance & Tax", "GST Documents", ["Company Name", "Filing Period", "Filing Type"]),
    ("Corporate", "Finance & Tax", "Other", ["Company Name", "Document Title"]),
    ("Corporate", "Finance & Tax", "Tax Documents", ["Company Name", "Assessment Year", "Tax Type"]),
    ("Corporate", "Intellectual Property", "Copyright Documents", ["Company Name", "Work Title", "Copyright Office"]),
    ("Corporate", "Intellectual Property", "IP Assignment Agreements", ["Company Name", "IP Description"]),
    ("Corporate", "Intellectual Property", "Other", ["Company Name", "IP Type"]),
    ("Corporate", "Intellectual Property", "Patent Documents", ["Company Name", "Invention Title", "Patent Office", "Validity"]),
    ("Corporate", "Intellectual Property", "Trademark Documents", ["Company Name", "Trademark Name", "Class", "Validity"]),
    ("Corporate", "Other", "Other", ["Company Name", "Document Title", "Issuing Authority"]),
    ("Legal", "Agreements", "Employment", ["Employer Name", "Designation"]),
    ("Legal", "Agreements", "MOUs", ["Purpose", "Duration/Term"]),
    ("Legal", "Agreements", "NDA", ["Duration/Term", "Subject Matter"]),
    ("Legal", "Agreements", "Other", ["Agreement Type"]),
    ("Legal", "Agreements", "Partnership", ["Profit Sharing Ratio"]),
    ("Legal", "Agreements", "Proprietorship", ["Business Name"]),
    ("Legal", "Agreements", "Sale", ["Subject Matter"]),
    ("Legal", "Client Management", "Client Profiles", ["Client Type (Individual/Corporate)"]),
    ("Legal", "Client Management", "Engagement Letters", ["Matter/Case Title", "Scope of Work"]),
    ("Legal", "Client Management", "Identity Proof", ["ID Type"]),
    ("Legal", "Client Management", "Other", ["Document Title"]),
    ("Legal", "Corporate Legal", "Corporate Legal Documents", ["Company Name", "Document Title", "Matter Type"]),
    ("Legal", "Corporate Legal", "Other", ["Company Name", "Document Title"]),
    ("Legal", "Court Documents", "Affidavits", ["Court Name", "Subject Matter"]),
    ("Legal", "Court Documents", "Evidence", ["Court Name", "Evidence Description"]),
    ("Legal", "Court Documents", "Orders", ["Court Name", "Order Summary"]),
    ("Legal", "Court Documents", "Other", ["Court Name", "Document Title"]),
    ("Legal", "Court Documents", "Petitions", ["Court Name", "Petition Type"]),
    ("Legal", "Litigation Calendar", "Litigation-related Calendar Documents", ["Court Name", "Hearing Type"]),
    ("Legal", "Litigation Calendar", "Other", ["Event Type"]),
    ("Legal", "Other", "Other", ["Document Title"]),
    ("Legal", "Patents", "Other", ["Issuing Authority"]),
    ("Legal", "Patents", "Patent Documents", ["Invention Title", "Patent Office", "Validity"]),
    ("Personal", "Education", "10th/12th Marksheet", ["School/Board Name", "Class/Standard", "Grade", "Year of Passing"]),
    ("Personal", "Education", "Bonafide Certificate", ["School/College Name", "Class/Course", "Academic Year", "Purpose"]),
    ("Personal", "Education", "Hall Ticket/Admit Card", ["Exam Name", "Exam Centre", "School/College Name"]),
    ("Personal", "Education", "Other", ["Institution Name", "Document Title"]),
    ("Personal", "Education", "School/College ID Card", ["School/College Name", "Class/Course", "Valid Until"]),
    ("Personal", "Education", "Semester Marksheet", ["University/Institute Name", "Course/Branch", "Semester"]),
    ("Personal", "Education", "Transfer Certificate", ["School/College Name", "Class Last Studied", "Reason for Leaving"]),
    ("Personal", "Financial", "Bank Account Documents", ["Bank Name", "Account Type", "Statement Period"]),
    ("Personal", "Financial", "Fixed Deposits", ["Bank Name", "Interest Rate", "Tenure"]),
    ("Personal", "Financial", "Form 16", ["Employer Name", "Assessment Year", "Financial Year"]),
    ("Personal", "Financial", "IT Returns", ["Assessment Year", "Financial Year", "ITR Form Type"]),
    ("Personal", "Financial", "Loan Documents", ["Loan Type", "Interest Rate", "Tenure"]),
    ("Personal", "Financial", "Mutual Funds", ["AMC/Fund House", "Scheme Name", "Plan/Option", "Units", "NAV"]),
    ("Personal", "Financial", "Other", ["Institution Name"]),
    ("Personal", "Gadgets/Appliances", "AMC", ["Product Name", "Brand", "Model Number", "Service Provider"]),
    ("Personal", "Gadgets/Appliances", "Laptop", ["Product Name", "Brand", "Model Number", "Product ID", "Dealer/Seller"]),
    ("Personal", "Gadgets/Appliances", "Other", ["Product Name", "Brand", "Model Number", "Warranty Expiry"]),
    ("Personal", "Gadgets/Appliances", "Purchase Invoice", ["Product Name", "Brand", "Model Number", "Quantity"]),
    ("Personal", "Gadgets/Appliances", "Refrigerator", ["Product Name", "Brand", "Model Number", "Product ID", "Dealer/Seller"]),
    ("Personal", "Gadgets/Appliances", "Robo Cleaner", ["Product Name", "Brand", "Model Number", "Product ID", "Dealer/Seller"]),
    ("Personal", "Gadgets/Appliances", "Service Record", ["Product Name", "Model Number", "Service Centre", "Issue/Complaint", "Service Performed", "Parts Replaced"]),
    ("Personal", "Gadgets/Appliances", "TV", ["Product Name", "Brand", "Model Number", "Product ID", "Dealer/Seller"]),
    ("Personal", "Gadgets/Appliances", "User Manual", ["Product Name", "Brand", "Model Number", "Product Type", "Manual Version/Edition"]),
    ("Personal", "Gadgets/Appliances", "Warranty", ["Brand", "Product Name", "Model Number", "Dealer/Service Provider"]),
    ("Personal", "Gadgets/Appliances", "Washing Machine", ["Product Name", "Brand", "Model Number", "Product ID", "Dealer/Seller"]),
    ("Personal", "Government Certificates", "Caste Certificate", ["Issuing Authority"]),
    ("Personal", "Government Certificates", "Domicile Certificate", ["Issuing Authority", "Purpose"]),
    ("Personal", "Government Certificates", "Income Certificate", ["Financial Year", "Issuing Authority", "Validity"]),
    ("Personal", "Government Certificates", "Non-Creamy Layer Certificate", ["Issuing Authority", "Validity"]),
    ("Personal", "Government Certificates", "Other", ["Certificate Title", "Issuing Authority"]),
    ("Personal", "Healthcare", "Blood Group Information", ["Hospital/Lab Name"]),
    ("Personal", "Healthcare", "Medical Reports", ["Hospital/Lab Name", "Report Type"]),
    ("Personal", "Healthcare", "Other", ["Document Type"]),
    ("Personal", "Healthcare", "Prescriptions", ["Hospital/Clinic", "Duration"]),
    ("Personal", "Healthcare", "Vaccinations", ["Dose Number", "Batch Number", "Vaccination Centre"]),
    ("Personal", "Identity & Legal", "Birth Certificate", ["Issuing Authority"]),
    ("Personal", "Identity & Legal", "Driving Licence", ["Vehicle Classes", "Issuing Authority/RTO"]),
    ("Personal", "Identity & Legal", "Marriage Certificate", ["Issuing Authority"]),
    ("Personal", "Identity & Legal", "Other", ["Document Title", "Issuing Authority"]),
    ("Personal", "Identity & Legal", "Passport", ["Place of Issue"]),
    ("Personal", "Identity & Legal", "Ration Card", ["Card Type", "Issuing Authority"]),
    ("Personal", "Insurance", "Health Insurance", ["Insurance Company", "Policy Type", "TPA/Network Details"]),
    ("Personal", "Insurance", "Home Insurance", ["Insurance Company", "Policy Type"]),
    ("Personal", "Insurance", "Life Insurance", ["Insurance Company", "Policy Type"]),
    ("Personal", "Insurance", "Other Insurance", ["Insurance Company", "Policy Type"]),
    ("Personal", "Insurance", "Travel Insurance", ["Insurance Company", "Policy Type", "Destination"]),
    ("Personal", "Insurance", "Vehicle Insurance", ["Insurance Company", "Policy Type", "Vehicle Make/Model"]),
    ("Personal", "Jewellery", "Hallmark Certificate", ["Jewellery Description", "Metal Type", "Purity", "Weight", "Jeweller Name"]),
    ("Personal", "Jewellery", "Insurance", ["Jewellery Description", "Insurance Company"]),
    ("Personal", "Jewellery", "Other", ["Jewellery Description", "Metal Type", "Purity", "Weight"]),
    ("Personal", "Jewellery", "Purchase Invoice", ["Jeweller Name", "Jewellery Description", "Metal Type", "Purity", "Gross/Net Weight", "Stone Weight"]),
    ("Personal", "Jewellery", "Valuation Certificate", ["Jewellery Description", "Metal Type", "Purity", "Gross/Net Weight", "Stone Details"]),
    ("Personal", "Jewellery", "Warranty", ["Jewellery Description", "Jeweller Name"]),
    ("Personal", "Property", "Lease Agreement", ["Lease Duration", "Notice Period"]),
    ("Personal", "Property", "Other", ["Registration Authority"]),
    ("Personal", "Property", "Property Tax", ["Tax Year"]),
    ("Personal", "Property", "Purchase Documents", ["Property Type", "Registration Authority"]),
    ("Personal", "Property", "Sale Deed", ["Property Description", "Registrar Office"]),
    ("Personal", "Travel", "Flight Ticket", ["Airline", "Flight Number", "Departure/Arrival Airport", "Class"]),
    ("Personal", "Travel", "Foreign Exchange Records", ["Exchange Provider/Bank", "Currency", "Exchange Rate"]),
    ("Personal", "Travel", "Hotel Booking", ["Hotel Name", "Number of Guests", "Room Type", "Number of Rooms"]),
    ("Personal", "Travel", "Other", ["Destination", "Provider"]),
    ("Personal", "Travel", "Travel Insurance", ["Insurance Company", "Destination"]),
    ("Personal", "Travel", "Visa", ["Visa Type", "Country", "Number of Entries", "Duration of Stay"]),
    ("Personal", "Utilities & Bills", "Broadband Bill", ["Billing Period", "Plan Name", "Provider Name"]),
    ("Personal", "Utilities & Bills", "Electricity Bill", ["Billing Period", "Units Consumed", "Provider Name"]),
    ("Personal", "Utilities & Bills", "Gas Bill", ["Billing Period", "Provider Name"]),
    ("Personal", "Utilities & Bills", "Mobile Bill", ["Billing Period", "Plan Name", "Provider Name"]),
    ("Personal", "Utilities & Bills", "Other Utility Bill", ["Billing Period", "Provider Name"]),
    ("Personal", "Utilities & Bills", "Water Bill", ["Billing Period", "Provider Name"]),
    ("Personal", "Vehicle", "Other", ["Vehicle Make/Model"]),
    ("Personal", "Vehicle", "PUC", ["Fuel Type", "Emission Values", "Testing Centre"]),
    ("Personal", "Vehicle", "Purchase Warranty", ["Vehicle Make/Model", "Dealer Name"]),
    ("Personal", "Vehicle", "RC Book", ["Vehicle Make/Model", "Vehicle Class", "Fuel Type", "RC Validity", "RTO"]),
    ("Personal", "Vehicle", "Road Tax", ["Tax Period", "RTO"]),
    ("Personal", "Vehicle", "Service History", ["Vehicle Make/Model", "Service Centre", "Odometer Reading", "Service Type", "Parts Replaced", "Next Service Due"]),
    ("Professional", "Awards & Recognition", "Awards", ["Award Title", "Awarding Organization", "Category/Field"]),
    ("Professional", "Awards & Recognition", "Certificates", ["Certificate Title", "Issuing Authority"]),
    ("Professional", "Awards & Recognition", "Other", ["Issuing Authority"]),
    ("Professional", "Awards & Recognition", "Recognition Documents", ["Recognition Title", "Issuing Organization"]),
    ("Professional", "Business", "Company Registration", ["Company Name", "Registrar of Companies"]),
    ("Professional", "Business", "GST Documents", ["Business Name", "Constitution of Business"]),
    ("Professional", "Business", "Licenses", ["Business Name", "Issuing Authority"]),
    ("Professional", "Business", "MSME", ["Business Name", "Business Category"]),
    ("Professional", "Business", "Other", ["Business Name", "Issuing Authority"]),
    ("Professional", "Business", "PAN/TAN", ["Entity Name"]),
    ("Professional", "Certification", "AI Course", ["Course Name", "Institution/Platform Name", "Grade/Score", "Duration"]),
    ("Professional", "Certification", "Degree", ["University/Institute Name", "Branch/Specialization", "Year of Passing"]),
    ("Professional", "Certification", "Memberships", ["Organization Name", "Membership Type", "Valid From/Until"]),
    ("Professional", "Certification", "Music", ["Course/Exam Name", "Institution Name", "Grade/Level"]),
    ("Professional", "Certification", "Others", ["Certificate Title", "Issuing Authority"]),
    ("Professional", "Certification", "Sports", ["Organizing Authority", "Position/Achievement"]),
    ("Professional", "Employment", "Appointment/Offer Letter", ["Employer/Company Name", "Designation", "Department", "Employment Type"]),
    ("Professional", "Employment", "Appraisal", ["Employer Name", "Designation", "Appraisal Period", "Rating"]),
    ("Professional", "Employment", "Experience Certificate", ["Employer Name", "Designation", "Department", "Duration of Employment"]),
    ("Professional", "Employment", "Other", ["Employer Name", "Designation"]),
    ("Professional", "Employment", "Promotion Letters", ["Employer Name", "Previous/New Designation"]),
    ("Professional", "Employment", "Relieving Letter", ["Employer Name", "Designation"]),
    ("Professional", "Employment", "Salary Slip", ["Employer Name", "Designation", "Pay Period"]),
    ("Professional", "IP (Intellectual Property)", "Copyright", ["Work Title", "Copyright Office"]),
    ("Professional", "IP (Intellectual Property)", "Granted Patent", ["Invention Title", "Patent Office", "Validity"]),
    ("Professional", "IP (Intellectual Property)", "Other", ["Issuing Authority"]),
    ("Professional", "IP (Intellectual Property)", "Patent Application", ["Invention Title", "Patent Office", "Status"]),
    ("Professional", "IP (Intellectual Property)", "Trademark", ["Trademark Name", "Class", "Validity"]),
]
def _low_tier_normalize(s):
    s = re.sub(r"\([^)]*\)", " ", s or "")
    s = re.sub(r"[^a-z0-9]+", " ", s.lower())
    return re.sub(r"\s+", " ", s).strip()
def _low_tier_strip_parenthetical(label):
    stripped = re.sub(r"\s*\([^)]*\)", "", label).strip()
    return stripped or label
LOW_TIER_EXACT_INDEX = {}
LOW_TIER_CATEGORY_INDEX = {}
for _ltd, _ltc, _ltdt, _ltfields in LOW_TIER_SAFE_FIELDS_TABLE:
    _lt_key = (_low_tier_normalize(_ltd), _low_tier_normalize(_ltc), _low_tier_normalize(_ltdt))
    LOW_TIER_EXACT_INDEX[_lt_key] = _ltfields
    _lt_cat_key = (_low_tier_normalize(_ltd), _low_tier_normalize(_ltc))
    LOW_TIER_CATEGORY_INDEX.setdefault(_lt_cat_key, []).append((_low_tier_normalize(_ltdt), _ltfields))
def get_low_tier_fields(domain, category, document_type, min_ratio=0.5):
    nd, nc, ndt = _low_tier_normalize(domain), _low_tier_normalize(category), _low_tier_normalize(document_type)
    exact = LOW_TIER_EXACT_INDEX.get((nd, nc, ndt))
    if exact is not None:
        return exact
    candidates = LOW_TIER_CATEGORY_INDEX.get((nd, nc), [])
    best_fields, best_ratio, other_fields = None, 0.0, None
    for cand_dt, cand_fields in candidates:
        if cand_dt == "other":
            other_fields = cand_fields
            continue
        ratio = SequenceMatcher(None, ndt, cand_dt).ratio()
        if ratio > best_ratio:
            best_ratio, best_fields = ratio, cand_fields
    if best_fields is not None and best_ratio >= min_ratio:
        return best_fields
    return other_fields or []
def _low_tier_match_or_learn(bucket_key, value, db, threshold=LOW_TIER_MATCH_THRESHOLD):
    bucket = db.setdefault(bucket_key, {})
    best_key, best_ratio = None, 0.0
    for existing in bucket:
        ratio = _fuzzy_ratio(existing, value)
        if ratio > best_ratio:
            best_key, best_ratio = existing, ratio
    if best_key is not None and best_ratio >= threshold:
        bucket[best_key] += 1
        return best_key, "matched_existing", round(best_ratio, 2)
    bucket[value] = bucket.get(value, 0) + 1
    return value, "auto_learned_new", 1.0
_VALUE_NOISE_EDGE_RE = re.compile(r'^[\s\*\|_~`^"\'.,;:!#@=<>\[\]{}\\/+]+|[\s\*\|_~`^"\'.,;:!#@=<>\[\]{}\\/+]+$')
def _clean_extracted_value(value):
    if value is None:
        return None
    cleaned = _VALUE_NOISE_EDGE_RE.sub("", value)
    cleaned = re.sub(r"\s{2,}", " ", cleaned).strip()
    return cleaned
def _low_tier_compulsory_lookup(field_name, compulsory_fields):
    if not compulsory_fields:
        return None
    target = _low_tier_normalize(field_name)
    required_fields = compulsory_fields.get("required_fields") or {}
    best_value, best_ratio = None, 0.0
    for label, info in required_fields.items():
        if not info.get("found"):
            continue
        ratio = SequenceMatcher(None, target, _low_tier_normalize(label)).ratio()
        if ratio > best_ratio:
            best_ratio, best_value = ratio, info["value"]
    if best_ratio >= 0.6:
        return best_value
    return None
def extract_low_tier_fields(raw_text, domain, category, document_type, compulsory_fields=None,
                            db=None, db_path=DEFAULT_LOW_TIER_DB_PATH,
                            threshold=LOW_TIER_MATCH_THRESHOLD, persist=True):
    allowed_fields = get_low_tier_fields(domain, category, document_type)
    if not allowed_fields:
        return {"document_type": document_type, "allowed_fields": [], "fields": {}}
    owns_db = db is None
    if owns_db:
        db = load_value_db(db_path)
    claimed_spans = []
    fields_out = {}
    for field_name in allowed_fields:
        clean_label = _low_tier_strip_parenthetical(field_name)
        pattern = _build_label_regex(clean_label)
        value = None
        matched_by = None
        for m in pattern.finditer(raw_text):
            span = m.span()
            if _span_overlaps(span, claimed_spans):
                continue
            candidate = m.group(1).strip(" \t:-")
            candidate = _clean_extracted_value(candidate)
            if candidate and _is_probably_a_label(candidate):
                candidate = None
            if candidate:
                claimed_spans.append(span)
                value = candidate
                matched_by = "label"
                break
        if value is None:
            nl_value, nl_span = _match_label_next_line(raw_text, clean_label, claimed_spans)
            if nl_value is not None:
                value = nl_value
                matched_by = "label_next_line"
        if value is None:
            value = _low_tier_compulsory_lookup(field_name, compulsory_fields)
            if value is not None:
                matched_by = "compulsory_fields_reuse"
        if value is None:
            field_type = _guess_field_type(clean_label)
            extractor = _FALLBACK_EXTRACTORS.get(field_type)
            if extractor is not None:
                fb_value, fb_span = extractor(raw_text, claimed_spans)
                fb_value = _clean_extracted_value(fb_value)
                if fb_value:
                    claimed_spans.append(fb_span)
                    value = fb_value
                    matched_by = "fallback_pattern"
        if value is None:
            fields_out[field_name] = {"value": None, "found": False, "match_status": None,
                                       "matched_by": None, "similarity": None}
            continue
        bucket_key = f"{document_type}::{field_name}"
        canonical_value, status, ratio = _low_tier_match_or_learn(bucket_key, value, db, threshold=threshold)
        fields_out[field_name] = {
            "value": canonical_value,
            "raw_extracted_value": value,
            "found": True,
            "match_status": status,
            "matched_by": matched_by,
            "similarity": ratio,
        }
    if owns_db and persist:
        save_value_db(db, db_path)
    return {"document_type": document_type, "allowed_fields": allowed_fields, "fields": fields_out}
def _print_low_tier_summary(low_tier_result, file=sys.stderr):
    if not low_tier_result or not low_tier_result.get("allowed_fields"):
        return
    print(f"\nLow-Tier Raw-Value Fields for '{low_tier_result['document_type']}':", file=file)
    for name, info in low_tier_result["fields"].items():
        if not info["found"]:
            print(f"  [ ] {name}: (not found)", file=file)
            continue
        tag = ""
        if info["match_status"] == "matched_existing":
            tag = f" [matched known value, similarity={info['similarity']}]"
        elif info["match_status"] == "auto_learned_new":
            tag = " [new value, auto-learned]"
        source = ""
        if info.get("matched_by") == "compulsory_fields_reuse":
            source = " (reused from compulsory fields)"
        elif info.get("matched_by") == "fallback_pattern":
            source = " (guessed, no label found)"
        print(f"  [x] {name}: {info['value']}{tag}{source}", file=file)
DB_HISTORY_MIN_OCCURRENCES = 2
def _compulsory_bucket_key(document_type, label):
    return f"{document_type}::{label}"
def _db_history_lookup(bucket_key, db, min_count=DB_HISTORY_MIN_OCCURRENCES):
    if not db:
        return None, 0
    bucket = db.get(bucket_key)
    if not bucket:
        return None, 0
    best_value, best_count = max(bucket.items(), key=lambda kv: kv[1])
    if best_count < min_count:
        return None, 0
    return best_value, best_count
def extract_compulsory_fields(raw_text, document_type, db=None):
    """Extract every compulsory field for `document_type`.

    When a field can't be read directly off the page (no label match, no
    fallback pattern), and a value-history DB is supplied, the most
    frequently-seen historical value for that exact (document_type, field)
    pair is used to auto-fill the field (`matched_by="db_history"`) instead
    of leaving it blank. Every field value that *is* read straight from the
    document gets written back into that same DB, so the pipeline keeps
    learning better auto-fill suggestions the more documents of a given
    type it sees.
    """
    path = DOC_TYPE_TO_PATH.get(document_type)
    if path is None:
        return None
    _, category = path
    key = f"{category}|{document_type}"
    if key not in REQUIRED_FIELDS:
        key = f"{category}|Other"
    required = REQUIRED_FIELDS.get(key)
    if not required:
        return None
    claimed_spans = []
    required_fields = {}
    found_count = 0
    fallback_count = 0
    db_filled_count = 0
    for label in required:
        pattern = _LABEL_PATTERNS.get(label) or _build_label_regex(label)
        value = None
        matched_by = None
        for m in pattern.finditer(raw_text):
            if _span_overlaps(m.span(), claimed_spans):
                continue
            candidate = m.group(1).strip(" \t:-")
            candidate = _clean_extracted_value(candidate)
            if candidate and _is_probably_a_label(candidate):
                candidate = None
            if candidate:
                claimed_spans.append(m.span())
                value = candidate
                matched_by = "label"
                break
        if value is None:
            value, span = _match_label_next_line(
                raw_text, label, claimed_spans, _LABEL_ONLY_PATTERNS.get(label))
            if value is not None:
                matched_by = "label_next_line"
        field_type = _guess_field_type(label)
        if value is None:
            extractor = _FALLBACK_EXTRACTORS.get(field_type)
            if extractor is not None:
                fb_value, fb_span = extractor(raw_text, claimed_spans)
                fb_value = _clean_extracted_value(fb_value)
                if fb_value:
                    claimed_spans.append(fb_span)
                    value = fb_value
                    matched_by = "fallback_pattern"
                    fallback_count += 1
        bucket_key = _compulsory_bucket_key(document_type, label)
        db_occurrences = None
        if value is not None:
            # A value was actually read off this document: reinforce (or
            # seed) the learning DB so future documents of this type can
            # be auto-filled/matched against it.
            if db is not None:
                canonical, _status, _ratio = _low_tier_match_or_learn(bucket_key, value, db)
                value = canonical
        elif db is not None:
            db_value, db_count = _db_history_lookup(bucket_key, db)
            if db_value is not None:
                value = db_value
                matched_by = "db_history"
                db_occurrences = db_count
                db_filled_count += 1
        required_fields[label] = {
            "value": value,
            "found": value is not None,
            "matched_by": matched_by,
            "source": "db_history" if matched_by == "db_history" else ("ocr" if value is not None else None),
            "db_occurrences": db_occurrences,
        }
        if value is not None:
            found_count += 1
    total = len(required)
    return {
        "document_type": document_type,
        "category": category,
        "required_fields": required_fields,
        "found_count": found_count,
        "total_count": total,
        "fallback_match_count": fallback_count,
        "db_filled_count": db_filled_count,
        "missing_fields": [lbl for lbl, v in required_fields.items() if not v["found"]],
        "completeness": round(found_count / total, 2) if total else None,
    }
def _print_fields_summary(fields, file=sys.stderr):
    print("Extracted Fields:", file=file)
    if not fields:
        print("  (none found)", file=file)
        return
    order = {"high": 0, "medium": 1, "low": 2}
    for name, data in sorted(fields.items(), key=lambda kv: order.get(kv[1]["confidence"], 1)):
        print(f"  [{data['confidence']:>6}] {name}: {', '.join(data['values'])}", file=file)
def _print_compulsory_fields_summary(compulsory, file=sys.stderr):
    if not compulsory:
        return
    print(f"\nCompulsory Fields for '{compulsory['document_type']}' "
          f"({compulsory['found_count']}/{compulsory['total_count']} found, "
          f"{compulsory['completeness']:.0%} complete, "
          f"{compulsory['fallback_match_count']} via label-less fallback, "
          f"{compulsory.get('db_filled_count', 0)} auto-filled from learned DB):", file=file)
    for label, info in compulsory["required_fields"].items():
        mark = "x" if info["found"] else " "
        value = info["value"] if info["found"] else "(not found)"
        if info["matched_by"] == "db_history":
            tag = f" [auto-filled from learned DB, seen {info.get('db_occurrences')}x before]"
        elif info["matched_by"] == "fallback_pattern":
            tag = " [guessed, no label found]"
        elif info["matched_by"] == "label_next_line":
            tag = " [label on own line, value taken from next line]"
        else:
            tag = ""
        print(f"  [{mark}] {label}: {value}{tag}", file=file)
def confirm_fields(document_type, fields, db_path=DEFAULT_LOW_TIER_DB_PATH, weight=3):
    """Record user-confirmed/corrected compulsory field values.

    A human confirming (or fixing) a value in the UI is a much stronger
    learning signal than a raw OCR read, so it's written into the same
    value-history DB used for `db_history` auto-fill, but with extra weight
    (`weight` occurrences added at once instead of 1), so confirmed values
    quickly outrank noisy OCR guesses and become the preferred auto-fill
    suggestion for that document type + field going forward.
    """
    db = load_value_db(db_path)
    updated = []
    for label, value in fields.items():
        value = _clean_extracted_value(str(value)) if value is not None else None
        if not value:
            continue
        bucket_key = _compulsory_bucket_key(document_type, label)
        bucket = db.setdefault(bucket_key, {})
        best_key, best_ratio = None, 0.0
        for existing in bucket:
            ratio = _fuzzy_ratio(existing, value)
            if ratio > best_ratio:
                best_key, best_ratio = existing, ratio
        if best_key is not None and best_ratio >= LOW_TIER_MATCH_THRESHOLD:
            bucket[best_key] += weight
            updated.append({"field": label, "value": best_key, "status": "reinforced"})
        else:
            bucket[value] = bucket.get(value, 0) + weight
            updated.append({"field": label, "value": value, "status": "learned_new"})
    save_value_db(db, db_path)
    return updated
def extract_tables(layout):
    return layout["table_rows"] if layout.get("table_detected") else []
COMMON_OCR_FIXES = [
    (re.compile(r"\bl\b"), "I"),
    (re.compile(r"0(?=[A-Za-z])"), "O"),
]
def validate_and_correct(recognition, low_conf_threshold=60):
    if recognition.get("source_type") == "native":
        has_text = bool(recognition.get("raw_text", "").strip())
        flags = [] if has_text else ["NO_TEXT_DETECTED"]
        return {"flags": flags, "low_confidence_words": [], "is_reliable": has_text}
    words = recognition["words"]
    low_conf_words = [w["text"] for w in words if w["conf"] < low_conf_threshold]
    flags = []
    if recognition["avg_confidence"] < low_conf_threshold:
        flags.append("LOW_OVERALL_CONFIDENCE")
    if not words:
        flags.append("NO_TEXT_DETECTED")
    if len(low_conf_words) > 0:
        flags.append(f"{len(low_conf_words)}_LOW_CONFIDENCE_WORDS")
    return {
        "flags": flags,
        "low_confidence_words": low_conf_words[:20],
        "is_reliable": recognition["avg_confidence"] >= low_conf_threshold and bool(words),
    }
_NOISE_LINE_RE = re.compile(r"^[\s\-_=~*.:|#]+$")
def clean_extracted_text(raw_text):
    if not raw_text:
        return ""
    text = re.sub(r"(\w)-\n(\w)", r"\1\2", raw_text)
    cleaned_lines = []
    for line in text.splitlines():
        line = line.strip()
        if not line or _NOISE_LINE_RE.match(line):
            continue
        cleaned_lines.append(re.sub(r"[ \t]{2,}", " ", line))
    return "\n".join(cleaned_lines)
def _clean_lines(lines):
    out = []
    for line in lines:
        line = line.strip()
        if not line or _NOISE_LINE_RE.match(line):
            continue
        out.append(re.sub(r"[ \t]{2,}", " ", line))
    return out
SUPPORTED_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".doc", ".docx", ".xls", ".xlsx"}
def _run_image_pipeline(path, lang="eng", min_conf=0, adaptive=True):
    stages = preprocess_image(path)
    recognition = detect_and_recognize(stages, lang=lang, min_conf=min_conf, adaptive=adaptive)
    recognition["source_type"] = "ocr"
    layout = analyze_layout(recognition["words"])
    return recognition, layout
def _rasterize_pdf_page(page, dpi=300):
    import fitz
    zoom = dpi / 72
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    img = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)
    color_flag = cv2.COLOR_RGBA2BGR if pix.n == 4 else cv2.COLOR_RGB2BGR
    return cv2.cvtColor(img, color_flag)
def _run_pdf_pipeline(path, lang="eng", min_conf=0, dpi=300, native_text_threshold=20, adaptive=True):
    try:
        import fitz
    except ImportError as e:
        raise ImportError(
            "Reading PDFs requires PyMuPDF: pip install pymupdf"
        ) from e
    doc = fitz.open(path)
    try:
        page_texts = [page.get_text("text").strip() for page in doc]
        avg_chars = mean(len(t) for t in page_texts) if page_texts else 0
        if avg_chars >= native_text_threshold:
            raw_text = "\n\n".join(t for t in page_texts if t)
            lines = [ln for t in page_texts for ln in t.split("\n") if ln.strip()]
            recognition = {"words": [], "raw_text": raw_text, "avg_confidence": 100.0,
                           "source_type": "native"}
            layout = {"lines": lines, "line_count": len(lines),
                      "table_detected": False, "table_rows": []}
            return recognition, layout
        all_words, raw_texts, confidences, all_lines, all_table_rows = [], [], [], [], []
        for page in doc:
            img = _rasterize_pdf_page(page, dpi=dpi)
            stages = preprocess_image_array(img)
            page_result = detect_and_recognize(stages, lang=lang, min_conf=min_conf, adaptive=adaptive)
            raw_texts.append(page_result["raw_text"])
            confidences.append(page_result["avg_confidence"])
            all_words.extend(page_result["words"])
            page_layout = analyze_layout(page_result["words"])
            all_lines.extend(page_layout["lines"])
            if page_layout["table_detected"]:
                all_table_rows.extend(page_layout["table_rows"])
        recognition = {
            "words": all_words,
            "raw_text": "\n\n".join(t for t in raw_texts if t),
            "avg_confidence": round(mean(confidences), 2) if confidences else 0.0,
            "source_type": "ocr",
        }
        layout = {
            "lines": all_lines,
            "line_count": len(all_lines),
            "table_detected": bool(all_table_rows),
            "table_rows": all_table_rows,
        }
        return recognition, layout
    finally:
        doc.close()
def _run_docx_pipeline(path):
    try:
        import docx
    except ImportError as e:
        raise ImportError(
            "Reading .docx files requires python-docx: pip install python-docx"
        ) from e
    document = docx.Document(path)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    table_rows = []
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                table_rows.append(cells)
    raw_text = "\n".join(paragraphs + [" | ".join(r) for r in table_rows])
    recognition = {"words": [], "raw_text": raw_text, "avg_confidence": 100.0,
                    "source_type": "native"}
    layout = {
        "lines": paragraphs,
        "line_count": len(paragraphs),
        "table_detected": bool(table_rows),
        "table_rows": table_rows,
    }
    return recognition, layout
def _run_xlsx_pipeline(path):
    try:
        import openpyxl
    except ImportError as e:
        raise ImportError(
            "Reading .xlsx files requires openpyxl: pip install openpyxl"
        ) from e
    wb = openpyxl.load_workbook(path, data_only=True)
    table_rows, lines = [], []
    for sheet in wb.worksheets:
        lines.append(f"[Sheet: {sheet.title}]")
        for row in sheet.iter_rows(values_only=True):
            cells = ["" if c is None else str(c) for c in row]
            if any(c.strip() for c in cells):
                table_rows.append(cells)
                lines.append(" | ".join(cells))
    raw_text = "\n".join(lines)
    recognition = {"words": [], "raw_text": raw_text, "avg_confidence": 100.0,
                    "source_type": "native"}
    layout = {
        "lines": lines,
        "line_count": len(lines),
        "table_detected": len(table_rows) >= 2,
        "table_rows": table_rows,
    }
    return recognition, layout
def _convert_with_libreoffice(path, target_ext):
    import subprocess
    import tempfile
    import shutil
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    if not soffice:
        raise RuntimeError(
            "Reading legacy .doc/.xls files requires LibreOffice ('soffice') "
            "for conversion. Install it (apt install libreoffice) or convert "
            "the file to .docx/.xlsx first."
        )
    workdir = tempfile.mkdtemp(prefix="ocr_pipeline_")
    subprocess.run(
        [soffice, "--headless", "--convert-to", target_ext, "--outdir", workdir, path],
        check=True, capture_output=True, timeout=120,
    )
    converted = os.path.join(workdir, f"{os.path.splitext(os.path.basename(path))[0]}.{target_ext}")
    if not os.path.exists(converted):
        raise RuntimeError(f"LibreOffice conversion failed for {path}")
    return converted
def _run_doc_pipeline(path):
    return _run_docx_pipeline(_convert_with_libreoffice(path, "docx"))
def _run_xls_pipeline(path):
    return _run_xlsx_pipeline(_convert_with_libreoffice(path, "xlsx"))
def run_ocr_pipeline(file_path, lang="eng", min_conf=0, model_path=DEFAULT_MODEL_PATH,
                     auto_learn_enabled=True, auto_learn_min_confidence=0.75,
                     auto_learn_retrain_every=25,
                     low_tier_enabled=True, low_tier_db_path=DEFAULT_LOW_TIER_DB_PATH,
                     low_tier_threshold=LOW_TIER_MATCH_THRESHOLD,
                     document_type_hint=None, adaptive_ocr=True, db_assist_enabled=True):
    ext = os.path.splitext(file_path)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. Supported: "
            f"{', '.join(sorted(e.lstrip('.') for e in SUPPORTED_EXTENSIONS))}"
        )
    if ext in (".jpg", ".jpeg", ".png"):
        recognition, layout = _run_image_pipeline(file_path, lang=lang, min_conf=min_conf, adaptive=adaptive_ocr)
    elif ext == ".pdf":
        recognition, layout = _run_pdf_pipeline(file_path, lang=lang, min_conf=min_conf, adaptive=adaptive_ocr)
    elif ext == ".docx":
        recognition, layout = _run_docx_pipeline(file_path)
    elif ext == ".doc":
        recognition, layout = _run_doc_pipeline(file_path)
    elif ext == ".xlsx":
        recognition, layout = _run_xlsx_pipeline(file_path)
    else:
        recognition, layout = _run_xls_pipeline(file_path)
    recognition["raw_text"] = clean_extracted_text(recognition["raw_text"])
    layout["lines"] = _clean_lines(layout["lines"])
    layout["line_count"] = len(layout["lines"])
    model = load_classifier(model_path)
    classification = classify_document(recognition["raw_text"], model=model)
    if document_type_hint and document_type_hint in DOC_TYPE_TO_PATH:
        hint_domain, hint_category = DOC_TYPE_TO_PATH[document_type_hint]
        classification = {
            "domain": hint_domain,
            "category": hint_category,
            "document_type": document_type_hint,
            "confidence": 1.0,
            "method": "hint_override",
        }
    fields = extract_fields(recognition["raw_text"])
    # Load the learned value-history DB once and share it across compulsory
    # field auto-fill and low-tier field matching, so both stages read/write
    # the same in-memory store and it's persisted a single time at the end.
    value_db = load_value_db(low_tier_db_path) if (low_tier_enabled or db_assist_enabled) else None
    compulsory_fields = extract_compulsory_fields(
        recognition["raw_text"], classification["document_type"],
        db=value_db if db_assist_enabled else None,
    )
    tables = extract_tables(layout)
    validation = validate_and_correct(recognition)
    low_tier_fields = {"document_type": classification["document_type"], "allowed_fields": [], "fields": {}}
    if low_tier_enabled:
        low_tier_fields = extract_low_tier_fields(
            recognition["raw_text"], classification["domain"], classification["category"],
            classification["document_type"], compulsory_fields=compulsory_fields,
            db=value_db, db_path=low_tier_db_path, threshold=low_tier_threshold, persist=False,
        )
    if value_db is not None:
        save_value_db(value_db, low_tier_db_path)
    learn_result = {"auto_logged": False, "retrained": False}
    if auto_learn_enabled and _ML_AVAILABLE:
        learn_result = auto_learn(
            recognition["raw_text"], classification, model_path=model_path,
            min_confidence=auto_learn_min_confidence, retrain_every=auto_learn_retrain_every,
        )
    result = {
        "source_file": os.path.basename(file_path),
        "file_type": ext.lstrip("."),
        "processed_at": datetime.now(timezone.utc).isoformat(),
        "ocr": {
            "raw_text": recognition["raw_text"],
            "average_confidence": recognition["avg_confidence"],
            "word_count": len(recognition["words"]),
            "source": recognition["source_type"],
            "ensemble_method": recognition.get("ensemble_method"),
            "models_compared": recognition.get("models_compared", []),
            "winning_model": recognition.get("winning_model"),
        },
        "layout": {
            "line_count": layout["line_count"],
            "lines": layout["lines"],
            "table_detected": layout["table_detected"],
        },
        "document_understanding": {
            "domain": classification["domain"],
            "category": classification["category"],
            "document_type": classification["document_type"],
            "classification_confidence": classification["confidence"],
            "classification_method": classification["method"],
            "fields": fields,
            "compulsory_fields": compulsory_fields,
            "low_tier_fields": low_tier_fields,
            "tables": tables,
        },
        "validation": validation,
        "auto_learning": learn_result,
    }
    return result
def _cli():
    parser = argparse.ArgumentParser(
        description="Document understanding pipeline with a trainable classifier. "
                     f"Supported file types: {', '.join(sorted(e.lstrip('.') for e in SUPPORTED_EXTENSIONS))}.")
    parser.add_argument("file", nargs="?", help="Path to a file to process")
    parser.add_argument("output", nargs="?", help="Where to save the resulting JSON")
    parser.add_argument("--train", action="store_true",
                          help="(Re)train the ML classifier on seed + logged feedback data")
    parser.add_argument("--correct", nargs=2, metavar=("RESULT_JSON", "DOCUMENT_TYPE"),
                          help="Log the correct label for a previous result's OCR text, "
                              "e.g. --correct out.json \"Electricity Bill\"")
    parser.add_argument("--model", default=DEFAULT_MODEL_PATH,
                          help=f"Model file path (default: {DEFAULT_MODEL_PATH})")
    parser.add_argument("--feedback", default=DEFAULT_FEEDBACK_PATH,
                          help=f"Feedback log path (default: {DEFAULT_FEEDBACK_PATH})")
    parser.add_argument("--list-types", action="store_true",
                          help="Print every valid document_type label and exit")
    parser.add_argument("--no-auto-learn", action="store_true",
                          help="Disable auto-learning for this run (see AUTO-LEARNING above)")
    parser.add_argument("--auto-learn-confidence", type=float, default=0.75,
                          help="Min hybrid-classification confidence to auto-learn from (default: 0.75)")
    parser.add_argument("--auto-learn-every", type=int, default=25,
                          help="Auto-retrain the model after this many auto-logged examples (default: 25)")
    parser.add_argument("--low-tier-db", default=DEFAULT_LOW_TIER_DB_PATH,
                          help=f"Low-tier raw-value store path (default: {DEFAULT_LOW_TIER_DB_PATH})")
    parser.add_argument("--no-low-tier", action="store_true",
                          help="Disable low-tier raw-value field extraction/storage for this run")
    parser.add_argument("--low-tier-threshold", type=float, default=LOW_TIER_MATCH_THRESHOLD,
                          help=f"Similarity threshold to match a new value against a stored one "
                              f"(default: {LOW_TIER_MATCH_THRESHOLD})")
    parser.add_argument("--list-low-tier-db", action="store_true",
                          help="Print the low-tier raw-value store and exit")
    parser.add_argument("--document-type-hint", default=None,
                          help="Skip text-based classification and force this exact document "
                               "type (must match a TAXONOMY leaf, e.g. 'Aadhaar Card') for "
                               "compulsory/low-tier field extraction")
    parser.add_argument("--thorough", action="store_true",
                          help="Disable the adaptive OCR fast path and always run the full "
                               "4-model ensemble, even on clean/confident scans (slower, "
                               "occasionally marginally more accurate)")
    parser.add_argument("--no-db-assist", action="store_true",
                          help="Disable DB-assisted auto-fill/auto-learn for compulsory fields")
    parser.add_argument("--confirm-fields", nargs=2, metavar=("DOCUMENT_TYPE", "FIELDS_JSON"),
                          help="Log user-confirmed compulsory field values (a JSON object "
                               "mapping field label -> value) as a high-weight learning "
                               "signal, e.g. --confirm-fields \"Aadhaar Card\" "
                               "'{\"Full Name\": \"Jane Doe\"}'")
    args = parser.parse_args()
    if args.list_types:
        for dt in ALL_DOC_TYPES:
            domain, category = DOC_TYPE_TO_PATH[dt]
            print(f"{domain} / {category} / {dt}")
        return
    if args.list_low_tier_db:
        db = load_value_db(args.low_tier_db)
        if not db:
            print(f"No entries yet in {args.low_tier_db}.")
            return
        for bucket_key, values in sorted(db.items()):
            print(f"\n{bucket_key} ({len(values)} value(s)):")
            for value, count in sorted(values.items(), key=lambda kv: -kv[1]):
                print(f"  [{count:>3}x] {value}")
        return
    if args.confirm_fields:
        document_type, fields_json = args.confirm_fields
        try:
            fields = json.loads(fields_json)
        except json.JSONDecodeError as e:
            print(f"Invalid FIELDS_JSON: {e}", file=sys.stderr)
            sys.exit(1)
        updated = confirm_fields(document_type, fields, db_path=args.low_tier_db)
        print(f"Learned {len(updated)} confirmed field value(s) for '{document_type}'.")
        for u in updated:
            print(f"  {u['field']}: {u['value']} ({u['status']})")
        return
    if args.train:
        stats = train_classifier(feedback_path=args.feedback, model_path=args.model)
        print(f"Trained classifier on {stats['seed_examples']} seed examples + "
              f"{stats['real_examples']} real feedback examples "
              f"({stats['classes']} document types). Saved to {stats['model_path']}.")
        return
    if args.correct:
        result_path, document_type = args.correct
        with open(result_path, "r", encoding="utf-8") as f:
            prior_result = json.load(f)
        raw_text = prior_result.get("ocr", {}).get("raw_text", "")
        if not raw_text:
            print(f"No ocr.raw_text found in {result_path}; nothing to log.")
            sys.exit(1)
        add_training_example(raw_text, document_type, feedback_path=args.feedback)
        print(f"Logged correction ('{document_type}') to {args.feedback}. "
              f"Run --train to fold it into the model.")
        return
    if not args.file:
        parser.print_help()
        sys.exit(1)
    result = run_ocr_pipeline(
        args.file, model_path=args.model,
        auto_learn_enabled=not args.no_auto_learn,
        auto_learn_min_confidence=args.auto_learn_confidence,
        auto_learn_retrain_every=args.auto_learn_every,
        low_tier_enabled=not args.no_low_tier,
        low_tier_db_path=args.low_tier_db,
        low_tier_threshold=args.low_tier_threshold,
        document_type_hint=args.document_type_hint,
        adaptive_ocr=not args.thorough,
        db_assist_enabled=not args.no_db_assist,
    )
    output_json = json.dumps(result, indent=2, ensure_ascii=False)
    if args.output:
        with open(args.output, "w", encoding="utf-8") as f:
            f.write(output_json)
        print(f"Saved structured OCR result to {args.output}")
    else:
        print(output_json)
    if result["auto_learning"]["auto_logged"]:
        msg = "Auto-logged this result as a new training example"
        if result["auto_learning"]["retrained"]:
            msg += " and retrained the classifier"
        print(f"{msg}.", file=sys.stderr)
    _print_ocr_model_summary(result["ocr"], file=sys.stderr)
    _print_fields_summary(result["document_understanding"]["fields"], file=sys.stderr)
    _print_compulsory_fields_summary(result["document_understanding"]["compulsory_fields"], file=sys.stderr)
    _print_low_tier_summary(result["document_understanding"]["low_tier_fields"], file=sys.stderr)
if __name__ == "__main__":
    _cli()
